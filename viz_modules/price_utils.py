#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Модуль работы с ценами:
- Загрузка прайса из XLSX
- Работа с базой цен SQLite
- Поиск цен по параметрам
"""
import math
import os
import re
import sqlite3

from core.project_paths import BASE_DIR, PRICE_DB_PATH, PRICE_XLSX_PATH
from core.price_db import length_m_to_price_length_dm, parse_plate_price_rows_from_xlsx

try:
    import pandas as pd
except Exception:
    pd = None

try:
    from docx import Document
except Exception:
    Document = None


def load_price_table_from_xlsx(path: str):
    """Загружает таблицу цен вида: ключ length_dm -> {6:price,8:price,10:price,12:price}."""
    table = {}
    if pd is None:
        return table

    candidate_paths = []
    if os.path.exists(path):
        candidate_paths = [path]
    else:
        search_dirs = [
            os.path.dirname(path) if os.path.dirname(path) else BASE_DIR,
            BASE_DIR,
            os.path.join(BASE_DIR, 'банк знаний')
        ]
        for d in search_dirs:
            if not os.path.isdir(d):
                continue
            for name in os.listdir(d):
                low = name.lower()
                if low.endswith('.xlsx') and ('нов' in low and 'цен' in low):
                    candidate_paths.append(os.path.join(d, name))

    if not candidate_paths:
        print('[ПРАЙС] Файл не найден. Искал около:', path)
        return table

    chosen = None
    for p in candidate_paths:
        try:
            pd.read_excel(p, sheet_name=None)
            chosen = p
            break
        except Exception:
            continue

    if chosen is None:
        print('[ПРАЙС] Не удалось открыть ни один XLSX из кандидатов:', candidate_paths)
        return table

    print('[ПРАЙС] Использую прайс-файл:', chosen)
    rows = parse_plate_price_rows_from_xlsx(chosen)
    for length_dm, load_code, price in rows:
        table.setdefault(int(length_dm), {})[int(load_code)] = float(price)
    print(f"[ПРАЙС] Считано позиций: {len({length_dm for length_dm, _, _ in rows})}")
    return table


def sync_price_xlsx_to_db(xlsx_path: str = PRICE_XLSX_PATH, db_path: str = PRICE_DB_PATH,
                          sheet_hint: str = '24.06.2024') -> int:
    """Заливает прайс из XLSX в SQLite."""
    if pd is None:
        return 0
    price_table = load_price_table_from_xlsx(xlsx_path)
    if not price_table:
        return 0
    rows = []
    for length_dm, loads in price_table.items():
        for load_code, price in loads.items():
            rows.append((int(length_dm), int(load_code), float(price)))

    conn = sqlite3.connect(db_path)
    try:
        cur = conn.cursor()
        cur.execute('CREATE TABLE IF NOT EXISTS prices (length_dm INTEGER, load_code INTEGER, price REAL, PRIMARY KEY(length_dm, load_code))')
        cur.executemany('INSERT OR REPLACE INTO prices (length_dm, load_code, price) VALUES (?,?,?)', rows)
        conn.commit()
        return len(rows)
    finally:
        conn.close()


def find_price_from_db(length_m: float, load_code: float | int = 8, db_path: str = PRICE_DB_PATH) -> float:
    """
    Ищет цену в БД с допуском ±1 дм.
    
    ВАЖНО: Для нагрузки 12.5 использует цену 12п (math.floor).
    """
    import math
    
    length_dm = int(round(length_m * 10))
    
    # КЛЮЧЕВОЕ ИЗМЕНЕНИЕ: округляем нагрузку вниз (12.5 → 12)
    load_code_for_db = int(math.floor(load_code)) if isinstance(load_code, (int, float)) else 8
    
    conn = sqlite3.connect(db_path)
    try:
        cur = conn.cursor()
        cur.execute('CREATE TABLE IF NOT EXISTS prices (length_dm INTEGER, load_code INTEGER, price REAL, PRIMARY KEY(length_dm, load_code))')
        cur.execute('SELECT price FROM prices WHERE length_dm=? AND load_code=?', (length_dm, load_code_for_db))
        row = cur.fetchone()
        if row:
            return float(row[0])
        cur.execute('SELECT price FROM prices WHERE ABS(length_dm-?)<=1 AND load_code=? ORDER BY ABS(length_dm-?) LIMIT 1', (length_dm, load_code_for_db, length_dm))
        row = cur.fetchone()
        return float(row[0]) if row else None
    finally:
        conn.close()


def find_price_for_plate(price_table: dict, length_m: float, load_code: int | float = 8) -> float | None:
    """Возвращает цену по длине и нагрузке. Для нагрузки 12,5 используется цена из колонки 12 (целая часть)."""
    key = int(round(length_m * 10))
    try:
        load_code_int = int(math.floor(load_code)) if load_code is not None else 8
    except (TypeError, ValueError):
        load_code_int = 8
    if key in price_table and load_code_int in price_table[key]:
        result = price_table[key][load_code_int]
    else:
        result = None
        for Ldm, loads in price_table.items():
            if abs(Ldm - key) <= 1 and load_code_int in loads:
                result = loads[load_code_int]
                break
    return result


def _find_price_for_plate_production_fallback(
    price_table: dict,
    length_m: float,
    load_code: int | float = 8,
) -> float | None:
    """XLSX fallback для производственной сметы: ключ длины через length_m_to_price_length_dm (ceil)."""
    length_dm_key = length_m_to_price_length_dm(length_m)
    try:
        load_code_int = int(math.floor(load_code)) if load_code is not None else 8
    except (TypeError, ValueError):
        load_code_int = 8
    if length_dm_key in price_table and load_code_int in price_table[length_dm_key]:
        return price_table[length_dm_key][load_code_int]
    for tbl_dm, loads in price_table.items():
        if abs(tbl_dm - length_dm_key) <= 1 and load_code_int in loads:
            return loads[load_code_int]
    return None


def load_cut_price_from_docx(path: str) -> float:
    """Пытается извлечь цену продольного реза из DOCX."""
    if Document is None or not os.path.exists(path):
        return 0.0
    try:
        doc = Document(path)
        text = '\n'.join([p.text for p in doc.paragraphs])
        candidates = []
        for m in re.finditer(r'(рез|вдоль|продоль)[^\d]{0,20}(\d+[\s\u202f\,\.]?\d*)', text.lower()):
            try:
                val = float(m.group(2).replace(' ', '').replace('\u202f', '').replace(',', '.'))
                candidates.append(val)
            except Exception:
                pass
        if candidates:
            return float(max(candidates))
        for table in doc.tables:
            for row in table.rows:
                row_text = ' '.join(c.text for c in row.cells).lower()
                if any(k in row_text for k in ['рез', 'вдоль', 'продоль']):
                    nums = re.findall(r'\d+[\s\u202f\,\.]?\d*', row_text)
                    for s in nums:
                        try:
                            val = float(s.replace(' ', '').replace('\u202f', '').replace(',', '.'))
                            candidates.append(val)
                        except Exception:
                            pass
        return float(max(candidates)) if candidates else 0.0
    except Exception:
        return 0.0

