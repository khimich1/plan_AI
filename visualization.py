#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Модуль визуализации и работы с ценами:
- Загрузка прайса из XLSX
- Работа с базой цен SQLite
- Построение сметы
- Визуализация раскладки плит
"""
import os
import re
import sqlite3
from datetime import datetime
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as patches
from matplotlib.lines import Line2D

import config_and_data as cfg
from optimization import OPT_PLAN, OPT_WIDTH_PRIORITY, optimize_cuts_pulp
from price_db import init_schema, import_from_xlsx, get_price

try:
    import pandas as pd
except Exception:
    pd = None

try:
    from docx import Document
except Exception:
    Document = None


# ==================== РАБОТА С ЦЕНАМИ ====================

def load_price_table_from_xlsx(path: str):
    """Загружает таблицу цен вида: ключ length_dm -> {6:price,8:price,10:price,12:price}."""
    table = {}
    if pd is None:
        return table
    
    candidate_paths = []
    if os.path.exists(path):
        candidate_paths = [path]
    else:
        search_dirs = [
            os.path.dirname(path) if os.path.dirname(path) else cfg.BASE_DIR,
            cfg.BASE_DIR,
            os.path.join(cfg.BASE_DIR, 'банк знаний')
        ]
        for d in search_dirs:
            if not os.path.isdir(d):
                continue
            for name in os.listdir(d):
                low = name.lower()
                if low.endswith('.xlsx') and ('нов' in low and 'цен' in low):
                    candidate_paths.append(os.path.join(d, name))
    
    if not candidate_paths:
        print('[ПРАЙС] Файл не найден. Искал около:', path)
        return table
    
    try:
        chosen = None
        for p in candidate_paths:
            try:
                all_sheets = pd.read_excel(p, sheet_name=None)
                chosen = p
                break
            except Exception:
                continue
        
        if chosen is None:
            print('[ПРАЙС] Не удалось открыть ни один XLSX из кандидатов:', candidate_paths)
            return {}
        else:
            print('[ПРАЙС] Использую прайс-файл:', chosen)
        
        preferred_sheet = None
        for s in list(all_sheets.keys()):
            if '24.06.2024' in str(s):
                preferred_sheet = s
                break
        sheets_iter = [preferred_sheet] if preferred_sheet in all_sheets else list(all_sheets.keys())
        
        for sheet_name in sheets_iter:
            df = all_sheets[sheet_name]
            try:
                print(f"[ПРАЙС] Лист: {sheet_name} | колонки: {[str(c) for c in df.columns]}")
            except Exception:
                pass
            
            name_col = next((c for c in df.columns if str(c).strip().lower() == 'наименование'), None) or \
                       next((c for c in df.columns if 'наимен' in str(c).lower()), None)
            if name_col is None:
                continue
            
            load_cols = {}
            header_map = {6: None, 8: None, 10: None, 12: None}
            for c in df.columns:
                cl = str(c).strip().lower()
                if cl == '6 нагрузка':
                    header_map[6] = c
                elif cl == '8 нагрузка':
                    header_map[8] = c
                elif cl == '10 нагрузка':
                    header_map[10] = c
                elif cl == '12 нагрузка':
                    header_map[12] = c
            
            for k,v in header_map.items():
                if v is not None:
                    load_cols[k] = v
            
            simple_price_col = next((c for c in df.columns if any(k in str(c).lower() for k in ['цен', 'руб', 'стоим'])), None)
            for c in df.columns:
                cl = str(c).lower()
                m = re.search(r'(\d+)\s*нагруз', cl)
                if m:
                    load_cols[int(m.group(1))] = c
                    continue
                m2 = re.search(r'(?:цен|руб|стоим)[^\d]{0,10}(6|8|10|12)\b', cl)
                if not m2:
                    m2 = re.search(r'\b(6|8|10|12)[^\d]{0,10}(?:цен|руб|стоим)', cl)
                if m2:
                    try:
                        load_cols[int(m2.group(1))] = c
                    except Exception:
                        pass
            
            found_rows = 0
            for _, row in df.iterrows():
                name = str(row.get(name_col, '')).strip()
                if not name:
                    continue
                L, _ = cfg.parse_name_to_sizes(name)
                if L is None:
                    continue
                key = int(round(L*10))
                price_by_load = {}
                if load_cols:
                    for load_code, col in load_cols.items():
                        try:
                            val = row[col]
                            if pd.notna(val):
                                price_by_load[load_code] = float(str(val).replace(' ', '').replace(',', '.'))
                        except Exception:
                            pass
                elif simple_price_col is not None:
                    try:
                        val = row[simple_price_col]
                        if pd.notna(val):
                            price_val = float(str(val).replace(' ', '').replace(',', '.'))
                            for load_code in [6, 8, 10, 12]:
                                price_by_load[load_code] = price_val
                    except Exception:
                        pass
                if price_by_load:
                    table[key] = price_by_load
                    found_rows += 1
            try:
                print(f"[ПРАЙС] Считано позиций на листе: {found_rows}")
            except Exception:
                pass
    except Exception:
        return {}
    return table


def sync_price_xlsx_to_db(xlsx_path: str = cfg.PRICE_XLSX_PATH, db_path: str = cfg.PRICE_DB_PATH,
                          sheet_hint: str = '24.06.2024') -> int:
    """Заливает прайс из XLSX в SQLite."""
    if pd is None:
        return 0
    price_table = load_price_table_from_xlsx(xlsx_path)
    if not price_table:
        return 0
    rows = []
    for length_dm, loads in price_table.items():
        for load_code, price in loads.items():
            rows.append((int(length_dm), int(load_code), float(price)))

    conn = sqlite3.connect(db_path)
    try:
        cur = conn.cursor()
        cur.execute('CREATE TABLE IF NOT EXISTS prices (length_dm INTEGER, load_code INTEGER, price REAL, PRIMARY KEY(length_dm, load_code))')
        cur.executemany('INSERT OR REPLACE INTO prices (length_dm, load_code, price) VALUES (?,?,?)', rows)
        conn.commit()
        return len(rows)
    finally:
        conn.close()


def find_price_from_db(length_m: float, load_code: int = 8, db_path: str = cfg.PRICE_DB_PATH) -> float:
    """Ищет цену в БД с допуском ±1 дм."""
    length_dm = int(round(length_m * 10))
    conn = sqlite3.connect(db_path)
    try:
        cur = conn.cursor()
        cur.execute('CREATE TABLE IF NOT EXISTS prices (length_dm INTEGER, load_code INTEGER, price REAL, PRIMARY KEY(length_dm, load_code))')
        cur.execute('SELECT price FROM prices WHERE length_dm=? AND load_code=?', (length_dm, load_code))
        row = cur.fetchone()
        if row:
            return float(row[0])
        cur.execute('SELECT price FROM prices WHERE ABS(length_dm-?)<=1 AND load_code=? ORDER BY ABS(length_dm-?) LIMIT 1', (length_dm, load_code, length_dm))
        row = cur.fetchone()
        return float(row[0]) if row else None
    finally:
        conn.close()


def find_price_for_plate(price_table: dict, length_m: float, load_code: int = 8) -> float:
    """Возвращает цену по длине и нагрузке."""
    key = int(round(length_m*10))
    if key in price_table and load_code in price_table[key]:
        return price_table[key][load_code]
    for Ldm, loads in price_table.items():
        if abs(Ldm - key) <= 1 and load_code in loads:
            return loads[load_code]
    return None


def load_cut_price_from_docx(path: str) -> float:
    """Пытается извлечь цену продольного реза из DOCX."""
    if Document is None or not os.path.exists(path):
        return 0.0
    try:
        doc = Document(path)
        text = '\n'.join([p.text for p in doc.paragraphs])
        candidates = []
        for m in re.finditer(r'(рез|вдоль|продоль)[^\d]{0,20}(\d+[\s\u202f\,\.]?\d*)', text.lower()):
            try:
                val = float(m.group(2).replace(' ', '').replace('\u202f', '').replace(',', '.'))
                candidates.append(val)
            except Exception:
                pass
        if candidates:
            return float(max(candidates))
        for table in doc.tables:
            for row in table.rows:
                row_text = ' '.join(c.text for c in row.cells).lower()
                if any(k in row_text for k in ['рез', 'вдоль', 'продоль']):
                    nums = re.findall(r'\d+[\s\u202f\,\.]?\d*', row_text)
                    for s in nums:
                        try:
                            val = float(s.replace(' ', '').replace('\u202f', '').replace(',', '.'))
                            candidates.append(val)
                        except Exception:
                            pass
        return float(max(candidates)) if candidates else 0.0
    except Exception:
        return 0.0


# ==================== ПОСТРОЕНИЕ ЗАКУПКИ И СМЕТЫ ====================

def build_procurement_items():
    """Формирует реальные позиции закупки из заказа пользователя."""
    items = []
    from collections import Counter
    
    # Приоритет 1: Используем реальный заказ из cfg.PLATES_*
    # Это то, что пользователь заказал - честная цена!
    all_plates = []
    for width_mm, plates_list in [
        (320, cfg.PLATES_0_32), (460, cfg.PLATES_0_46), (700, cfg.PLATES_0_70),
        (720, cfg.PLATES_0_72), (860, cfg.PLATES_0_86), (880, cfg.PLATES_0_88),
        (740, cfg.PLATES_0_74), (480, cfg.PLATES_0_48), (500, cfg.PLATES_0_50),
        (340, cfg.PLATES_0_34), (1080, cfg.PLATES_1_08), (1200, cfg.PLATES_1_2),
        (1000, cfg.PLATES_1_0)
    ]:
        if plates_list:
            length_counts = Counter(plates_list)
            for length, qty in length_counts.items():
                all_plates.append({
                    'length': length,
                    'width': width_mm / 1000.0,  # в метрах
                    'qty': qty
                })
    
    if all_plates:
        # Есть реальный заказ - используем его
        for plate in all_plates:
            # Определяем количество резов (примерная оценка)
            width_m = plate['width']
            
            # Продольные резы: если ширина < 1.2м, значит был рез
            long_cuts = 1 if width_m < 1.15 else 0
            
            # Поперечные резы: пока 0, они учтены в оптимизации
            trans_cuts = 0
            
            items.append({
                'length': plate['length'],
                'width': width_m,
                'qty': plate['qty'],
                'long_cuts': long_cuts,
                'trans_cuts': trans_cuts
            })
        
        return items
    
    # Приоритет 2: Используем старый OPT_PLAN (если нет заказа)
    if OPT_PLAN and OPT_PLAN.get('actions'):
        for act in OPT_PLAN['actions']:
            src_type, W1, W2, L, qty, lc, tc = act
            W1_m = W1 / 1000.0; W2_m = W2 / 1000.0 if W2 else 0
            if src_type == 'split':
                items.append({'length': round(L, 2), 'width': 1.2, 'qty': qty, 'long_cuts': lc, 'trans_cuts': tc, 'purpose': 'split_source'})
            elif src_type == 'narrow':
                items.append({'length': round(L, 2), 'width': W2_m, 'qty': qty, 'long_cuts': lc, 'trans_cuts': tc, 'purpose': 'narrow_source'})
            elif src_type == 'solid':
                items.append({'length': round(L, 2), 'width': W1_m, 'qty': qty, 'long_cuts': lc, 'trans_cuts': tc, 'purpose': 'solid'})
        agg = {}
        for it in items:
            key = (it['length'], it['width'], it['long_cuts'], it['trans_cuts'])
            agg[key] = agg.get(key, 0) + it['qty']
        result = []
        for (L, W, long_cuts, trans_cuts), qty in sorted(agg.items(), key=lambda x: (x[0][1], x[0][0])):
            result.append({'length': L, 'width': W, 'qty': qty, 'long_cuts': long_cuts, 'trans_cuts': trans_cuts})
        return result
    
    # Fallback: старая логика
    def mismatch_count(main_list, pair_demand):
        if not main_list or not pair_demand:
            return 0
        a = sorted(round(x, 2) for x in main_list)
        b = sorted(round(x, 2) for x in pair_demand)
        i = j = matches = 0
        while i < len(a) and j < len(b):
            if abs(a[i] - b[j]) <= 0.05:
                matches += 1; i += 1; j += 1
            elif a[i] < b[j]:
                i += 1
            else:
                j += 1
        return max(0, min(len(main_list), len(pair_demand)) - matches)

    pair_plan = {
        '0.32': mismatch_count(cfg.PLATES_0_32, cfg.PLATES_0_88),
        '0.46': mismatch_count(cfg.PLATES_0_46, cfg.PLATES_0_74),
        '0.72': mismatch_count(cfg.PLATES_0_72, cfg.PLATES_0_48),
        '0.70': mismatch_count(cfg.PLATES_0_70, cfg.PLATES_0_50),
        '0.86': mismatch_count(cfg.PLATES_0_86, cfg.PLATES_0_34),
    }
    
    for L in cfg.PLATES_1_2:
        items.append({'length': round(L, 1), 'width': 1.2, 'qty': 1, 'long_cuts': 0, 'trans_cuts': 0, 'purpose': 'as_is'})
    for L in cfg.PLATES_1_5_TO_1_2:
        items.append({'length': round(L, 1), 'width': 1.2, 'qty': 1, 'long_cuts': 0, 'trans_cuts': 0, 'purpose': 'to_1_2_main'})
        items.append({'length': round(L, 1), 'width': 0.3, 'qty': 1, 'long_cuts': 1, 'trans_cuts': 0, 'purpose': 'to_1_2_strip'})
    for L in cfg.PLATES_1_0:
        items.append({'length': round(L, 1), 'width': 1.0, 'qty': 1, 'long_cuts': 1, 'trans_cuts': 0, 'purpose': 'to_1_0_main'})
        items.append({'length': round(L, 1), 'width': 0.2, 'qty': 1, 'long_cuts': 1, 'trans_cuts': 0, 'purpose': 'to_1_0_strip'})
    for L in cfg.PLATES_1_08:
        items.append({'length': round(L, 1), 'width': 1.08, 'qty': 1, 'long_cuts': 1, 'trans_cuts': 0, 'purpose': 'to_1_08_main'})
        items.append({'length': round(L, 1), 'width': 0.12, 'qty': 1, 'long_cuts': 1, 'trans_cuts': 0, 'purpose': 'to_1_08_strip'})
    for L in cfg.PLATES_0_46:
        items.append({'length': round(L, 1), 'width': 0.46, 'qty': 1, 'long_cuts': 1, 'trans_cuts': 0, 'purpose': 'to_0_46_main'})
        items.append({'length': round(L, 1), 'width': 0.74, 'qty': 1, 'long_cuts': 1, 'trans_cuts': 0, 'purpose': 'to_0_46_strip'})
    
    mismatch = pair_plan['0.32']
    for idx, L in enumerate(cfg.PLATES_0_32):
        items.append({'length': round(L, 1), 'width': 0.32, 'qty': 1, 'long_cuts': 1, 'trans_cuts': 0, 'purpose': 'to_0_32_main'})
        trans = 1 if idx < mismatch else 0
        items.append({'length': round(L, 1), 'width': 0.88, 'qty': 1, 'long_cuts': 1, 'trans_cuts': trans, 'purpose': 'to_0_32_strip'})
    
    mismatch = pair_plan['0.72']
    for idx, L in enumerate(cfg.PLATES_0_72):
        items.append({'length': round(L, 1), 'width': 0.72, 'qty': 1, 'long_cuts': 1, 'trans_cuts': 0, 'purpose': 'to_0_72_main'})
        trans = 1 if idx < mismatch else 0
        items.append({'length': round(L, 1), 'width': 0.48, 'qty': 1, 'long_cuts': 1, 'trans_cuts': trans, 'purpose': 'to_0_72_strip'})
    
    mismatch = pair_plan['0.70']
    for idx, L in enumerate(cfg.PLATES_0_70):
        items.append({'length': round(L, 1), 'width': 0.70, 'qty': 1, 'long_cuts': 1, 'trans_cuts': 0, 'purpose': 'to_0_70_main'})
        trans = 1 if idx < mismatch else 0
        items.append({'length': round(L, 1), 'width': 0.50, 'qty': 1, 'long_cuts': 1, 'trans_cuts': trans, 'purpose': 'to_0_70_strip'})
    
    mismatch = pair_plan['0.86']
    for idx, L in enumerate(cfg.PLATES_0_86):
        items.append({'length': round(L, 1), 'width': 0.86, 'qty': 1, 'long_cuts': 1, 'trans_cuts': 0, 'purpose': 'to_0_86_main'})
        trans = 1 if idx < mismatch else 0
        items.append({'length': round(L, 1), 'width': 0.34, 'qty': 1, 'long_cuts': 1, 'trans_cuts': trans, 'purpose': 'to_0_86_strip'})
    
    agg = {}
    for it in items:
        key = (it['length'], it['width'], it['long_cuts'], it['trans_cuts'])
        agg[key] = agg.get(key, 0) + it['qty']
    result = []
    for (L, W, long_cuts, trans_cuts), qty in sorted(agg.items(), key=lambda x: (x[0][1], x[0][0])):
        result.append({'length': L, 'width': W, 'qty': qty, 'long_cuts': long_cuts, 'trans_cuts': trans_cuts})
    return result


def build_price_rows(price_table: dict, reinforcement_code: int = 8):
    """Формирует строки сметы."""
    items = build_procurement_items()
    rows = []
    total = 0.0
    idx = 1
    for it in items:
        L, W, qty = it['length'], it['width'], it['qty']
        long_cuts, trans_cuts = it['long_cuts'], it['trans_cuts']
        name = cfg.make_plate_name(L, W)
        
        load_code = 6 if W < 1.0 else reinforcement_code
        db_price = get_price(L, load_code, cfg.PRICE_DB_PATH)
        base_price_1_2m = db_price if db_price is not None else (find_price_for_plate(price_table, L, load_code) or 0.0)
        
        if base_price_1_2m > 0:
            width_factor = W / 1.2
            base_price = base_price_1_2m * width_factor
        else:
            base_price = 0.0
        
        cuts_cost = long_cuts * (cfg.LONG_CUT_PRICE_PER_M * L) + trans_cuts * cfg.TRANSVERSE_CUT_PRICE
        unit_price = base_price + cuts_cost
        weight = cfg.approximate_weight_kg(L, W)
        row_sum = unit_price * qty
        total += row_sum

        metadata = []
        if hasattr(cfg, 'consume_plate_metadata'):
            try:
                metadata = cfg.consume_plate_metadata(L, int(round(W * 1000)), qty)
            except Exception:
                metadata = []
        weeks = [m.get('forming_week') for m in metadata if m.get('forming_week') not in (None, '')]
        week_str = ", ".join(str(w) for w in sorted(set(weeks))) if weeks else ''
        contractors = [m.get('contractor') for m in metadata if m.get('contractor')]
        contractor_str = ", ".join(sorted(set(contractors))) if contractors else ''

        rows.append([
            idx,
            name,
            qty,
            'шт',
            week_str or '—',
            contractor_str or '—',
            f'{weight:.0f}',
            f'{unit_price:,.2f}'.replace(',', ' ').replace('.', ','),
            f'{row_sum:,.2f}'.replace(',', ' ').replace('.', ',')
        ])
        idx += 1
    return rows, total


# ==================== ПОСТРОЕНИЕ ПОСЛЕДОВАТЕЛЬНОСТИ ====================

def build_layout_sequence():
    """Формирует последовательность сегментов вдоль дорожки."""
    from optimization import OPT_CASCADING_PLAN
    sequence = []

    def plate_label(L: float, W: float) -> str:
        Ldm = int(round(L * 10))
        Wdm_val = round(W * 10, 1)
        if abs(Wdm_val - int(Wdm_val)) < 1e-6:
            Wdm = str(int(Wdm_val))
        else:
            Wdm = str(Wdm_val).replace('.', ',')
        return f'ПБ {Ldm}-{Wdm}-8п'
    
    # Приоритет 1: OPT_CASCADING_PLAN (новая каскадная оптимизация)
    print(f"[VISUAL] Проверяем OPT_CASCADING_PLAN: {OPT_CASCADING_PLAN is not None}")
    if OPT_CASCADING_PLAN and OPT_CASCADING_PLAN.get('primary_cuts'):
        print("[VISUAL] OK: Используем каскадную оптимизацию для визуализации")
        print(f"[VISUAL] Первичных резов: {len(OPT_CASCADING_PLAN.get('primary_cuts', []))}")
        print(f"[VISUAL] Вторичных резов: {len(OPT_CASCADING_PLAN.get('secondary_cuts', []))}")
        
        # Проверяем, есть ли 2D данные (plate_assignments)
        use_2d_data = 'plate_assignments' in OPT_CASCADING_PLAN and OPT_CASCADING_PLAN['plate_assignments']
        
        if use_2d_data:
            print("[VISUAL] ТОЧНО: Используем 2D данные с точными длинами")
        else:
            print("[VISUAL] ВНИМАНИЕ: 2D данных нет, используем приближение")
            
            # Собираем все плиты с их длинами из cfg
            all_plates_with_lengths = []
            for plates, width_mm in [
                (cfg.PLATES_1_2, 1200), (cfg.PLATES_1_08, 1080),
                (cfg.PLATES_0_32, 320), (cfg.PLATES_0_46, 460), (cfg.PLATES_0_70, 700),
                (cfg.PLATES_0_72, 720), (cfg.PLATES_0_86, 860), (cfg.PLATES_0_88, 880),
                (cfg.PLATES_0_74, 740), (cfg.PLATES_0_48, 480), (cfg.PLATES_0_50, 500),
                (cfg.PLATES_0_34, 340)
            ]:
                for length in plates:
                    all_plates_with_lengths.append({'length': length, 'width': width_mm})
            
            # Сортируем по ширине для соответствия с оптимизацией
            all_plates_with_lengths.sort(key=lambda x: (-x['width'], -x['length']))
        
        # Создаём карту поперечных резов: {(length, width): {target_length, remainder}}
        transverse_cut_map = {}
        if OPT_CASCADING_PLAN.get('transverse_cuts'):
            for tcut in OPT_CASCADING_PLAN['transverse_cuts']:
                key = (tcut['source_length'], tcut['source_width'])
                transverse_cut_map[key] = {
                    'target_length': tcut['target_length'],
                    'remainder': tcut['remainder']
                }
            print(f"[VISUAL] Найдено {len(transverse_cut_map)} типов поперечных резов: {list(transverse_cut_map.keys())}")
        
        # СТАРАЯ ЛОГИКА (если нет plate_assignments_with_transverse)
        # Создаём карту вторичных резов: {(source_length, остаток_мм): {'pattern': [...], 'qty': N}}
        secondary_cuts_info = {}
        if OPT_CASCADING_PLAN.get('secondary_cuts'):
            for sec_cut in OPT_CASCADING_PLAN['secondary_cuts']:
                source_mm = sec_cut['source']
                pieces = sec_cut.get('pieces', 1)
                cuts_list = sec_cut.get('cuts', [])
                qty = sec_cut['qty']  # Сколько остатков режется вторично
                
                # ВАЖНО: Получаем ИСХОДНЫЕ длины остатков (ДО поперечного реза!)
                source_lengths_list = sec_cut.get('source_lengths', [])
                # Результирующие длины (ПОСЛЕ поперечного реза)
                target_lengths_list = sec_cut.get('lengths', [])
                
                # Создаём шаблон вторичных резов для ОДНОГО остатка
                pattern = []
                if cuts_list:
                    target_width_mm = cuts_list[0]
                    # Для множественной резки (pieces >= 2) создаём несколько сегментов
                    # Для сужения (pieces == 1) создаём один сегмент
                    for _ in range(pieces):
                        pattern.append({
                            'width': target_width_mm / 1000.0,
                            'width_mm': target_width_mm,  # Ширина РЕЗУЛЬТАТА вторичного реза
                            'source_width_mm': source_mm,  # Ширина ОСТАТКА (для правильной метки)
                            'label': None,  # Метка будет создана позже с реальной длиной плиты
                            'target_length': target_lengths_list[0] if target_lengths_list else None  # Для поперечных резов
                        })
                
                # ИСПРАВЛЕНИЕ: Создаём запись для КАЖДОЙ ИСХОДНОЙ длины отдельно
                for i in range(qty):
                    # Используем ИСХОДНУЮ длину остатка (ДО поперечного реза!)
                    source_length = source_lengths_list[i] if i < len(source_lengths_list) else 6.0
                    key = (source_length, source_mm)  # Ключ теперь (ИСХОДНАЯ длина, ширина)!
                    
                    if key not in secondary_cuts_info:
                        secondary_cuts_info[key] = {
                            'pattern': pattern,
                            'qty': 0,
                            'used': 0
                        }
                    secondary_cuts_info[key]['qty'] += 1
        
        print(f"[VISUAL] Создано {len(secondary_cuts_info)} вариантов вторичных резов:")
        for (src_len, src_w), info in secondary_cuts_info.items():
            pattern_desc = ", ".join([f"{c['width_mm']}мм" for c in info['pattern']])
            print(f"  Остаток {src_len}м x {src_w}мм: {info['qty']} шт -> [{pattern_desc}]")
        
        # 1. Первичные резы с вторичными резами внутри остатков
        for cut in OPT_CASCADING_PLAN.get('primary_cuts', []):
            width_mm = cut['width']
            rest_mm = cut['rest']
            qty = cut['qty']
            
            # Получаем длины для этих плит
            if use_2d_data and 'lengths' in cut:
                # Используем точные длины из 2D оптимизации
                lengths_for_cut = cut['lengths']
                print(f"[VISUAL] Первичный рез {width_mm}мм: используем точные длины {lengths_for_cut}")
            else:
                # Используем приближение из all_plates_with_lengths
                matching_plates = [p for p in all_plates_with_lengths if p['width'] == width_mm]
                lengths_for_cut = [p['length'] for p in matching_plates[:qty]]
                # Если не хватает, дополняем средней длиной
                while len(lengths_for_cut) < qty:
                    lengths_for_cut.append(6.0 if not matching_plates else matching_plates[0]['length'])
            
            for i in range(qty):
                # Берём длину для этой плиты
                length = lengths_for_cut[i] if i < len(lengths_for_cut) else 6.0
                
                # ИСПРАВЛЕНИЕ: Проверяем вторичные резы для КОНКРЕТНОГО остатка (длина + ширина)
                sec_info = secondary_cuts_info.get((length, rest_mm))
                
                # ОТЛАДКА: Выводим информацию о поиске вторичных резов
                if rest_mm > 0:
                    print(f"[VISUAL] Ищем вторичные резы для остатка {length}м x {rest_mm}мм: {'НАЙДЕНО' if sec_info else 'НЕ НАЙДЕНО'}")
                
                # Проверяем, есть ли поперечный рез для этой плиты
                transverse_cut_info = transverse_cut_map.get((length, width_mm))
                
                if transverse_cut_info:
                    # Эта плита режется поперёк - добавляем с mode='transverse'
                    width_m = width_mm / 1000.0
                    target_length = transverse_cut_info['target_length']
                    remainder = transverse_cut_info['remainder']
                    
                    sequence.append({
                        'length': length,  # Исходная длина плиты
                        'mode': 'transverse',
                        'target_length': target_length,
                        'remainder': remainder,
                        'width': width_m,
                        'label_target': plate_label(target_length, width_m),
                        'label_remainder': f'Остаток {remainder:.2f}м'.replace('.', ',') if remainder > 0.1 else ''
                    })
                    print(f"[VISUAL] Плита с поперечным резом: {length}м x {width_mm}мм -> {target_length}м (остаток {remainder:.2f}м)")
                else:
                    # Обычная плита без поперечного реза
                    main_w = width_mm / 1000.0
                    rest_w = rest_mm / 1000.0
                    
                    # Специальная обработка для плит БЕЗ реза (rest = 0)
                    if rest_mm == 0:
                        sequence.append({
                            'length': length,
                            'mode': 'solid',
                            'label': plate_label(length, main_w)
                        })
                    else:
                        # Плиты С резом
                        # Проверяем, нужны ли вторичные резы для ЭТОЙ плиты
                        secondary_cuts_for_plate = None
                        if sec_info and sec_info['used'] < sec_info['qty']:
                            # Эта плита получает вторичные резы
                            # Создаём копию шаблона с правильными метками (с реальной длиной плиты)
                            secondary_cuts_for_plate = []
                            for sec_cut_template in sec_info['pattern']:
                                sec_width = sec_cut_template['width']
                                sec_width_mm = sec_cut_template['width_mm']
                                
                                # ВАЖНО: Проверяем поперечные резы для ВТОРИЧНЫХ плит!
                                sec_transverse = transverse_cut_map.get((length, sec_width_mm))
                                
                                if sec_transverse:
                                    # Вторичная плита с поперечным резом
                                    secondary_cuts_for_plate.append({
                                        'width': sec_width,
                                        'label': f'[2] {plate_label(sec_transverse["target_length"], sec_width)}',
                                        'transverse_cut': True,
                                        'target_length': sec_transverse['target_length'],
                                        'remainder': sec_transverse['remainder']
                                    })
                                    print(f"[VISUAL] Вторичный рез С поперечным: {length}м x {sec_width_mm}мм -> {sec_transverse['target_length']}м")
                                else:
                                    # Вторичная плита (может быть с поперечным резом!)
                                    # Проверяем, есть ли target_length в шаблоне (для transverse-резов)
                                    target_length = sec_cut_template.get('target_length')
                                    
                                    if target_length:
                                        # Это вторичный рез типа 'transverse' (поперечный + продольный)
                                        # Метка показывает результат ОБОИХ резов
                                        secondary_cuts_for_plate.append({
                                            'width': sec_width,
                                            'label': f'О {plate_label(target_length, sec_width)}',  # О = Остаток
                                            'has_transverse': True,  # Флаг для отрисовки красной линии
                                            'target_length': target_length  # Длина результата (для правильной отрисовки)
                                        })
                                    else:
                                        # Обычный вторичный рез (только продольный)
                                        # Используем ширину ОСТАТКА для подписи, а не результата
                                        source_width = sec_cut_template.get('source_width_mm', sec_width_mm * 1000) / 1000.0
                                        secondary_cuts_for_plate.append({
                                            'width': sec_width,
                                            'label': f'О {plate_label(length, source_width)}'  # О = Остаток
                                        })
                            sec_info['used'] += 1
                        
                        sequence.append({
                            'length': length,
                            'mode': 'split',
                            'main_w': main_w,
                            'rest_w': rest_w,
                            'label_main': plate_label(length, main_w),
                            'label_rest': f'+{rest_w:.2f}'.replace('.', ',') if not secondary_cuts_for_plate else None,
                            'secondary_cuts': secondary_cuts_for_plate
                        })
        
        if sequence:
            return sequence
    else:
        print("[VISUAL] ВНИМАНИЕ: OPT_CASCADING_PLAN не найден или пуст, используем старый метод")
    
    # Приоритет 2: OPT_PLAN (старая оптимизация)
    if OPT_PLAN and OPT_PLAN.get('actions'):
        for act in OPT_PLAN['actions']:
            src_type, W1, W2, L, qty, lc, tc = act
            W1_m = W1 / 1000.0; W2_m = W2 / 1000.0 if W2 else 0
            for _ in range(qty):
                if src_type == 'solid':
                    sequence.append({'length': L, 'mode': 'solid', 'label': plate_label(L, W1_m)})
                elif src_type == 'split':
                    rest_w = W2_m if W2_m < W1_m else (1.2 - W1_m)
                    rest_label = f'+{rest_w:.2f}'.replace('.', ',')
                    sequence.append({'length': L, 'mode': 'split', 'main_w': W1_m, 'rest_w': rest_w,
                                     'label_main': plate_label(L, W1_m), 'label_rest': rest_label})
                elif src_type == 'narrow':
                    delta = abs(W2_m - W1_m) if W2_m else 0
                    rest_label = f'-{delta:.2f}'.replace('.', ',') if delta > 0.001 else ''
                    sequence.append({'length': L, 'mode': 'split', 'main_w': W1_m, 'rest_w': delta,
                                     'label_main': plate_label(L, W1_m), 'label_rest': rest_label})
        return sequence
    
    # Приоритет 3: Fallback на старую логику
    for L in cfg.PLATES_1_2:
        sequence.append({'length': L, 'mode': 'solid', 'label': plate_label(L, 1.2)})
    for L in cfg.PLATES_1_5_TO_1_2:
        sequence.append({'length': L, 'mode': 'solid', 'label': plate_label(L, 1.2)})
    for L in cfg.PLATES_1_0:
        sequence.append({'length': L, 'mode': 'split', 'main_w': 1.0, 'rest_w': 0.2,
                         'label_main': plate_label(L, 1.0), 'label_rest': '+0,2'})
    for L in cfg.PLATES_1_08:
        sequence.append({'length': L, 'mode': 'split', 'main_w': 1.08, 'rest_w': 0.12,
                         'label_main': plate_label(L, 1.08), 'label_rest': '+0,12'})
    
    groups_map = {
        '0_32': (cfg.PLATES_0_32, 0.32, 0.88, '+0,88'),
        '0_46': (cfg.PLATES_0_46, 0.46, 0.74, '+0,74'),
        '0_70': (cfg.PLATES_0_70, 0.70, 0.50, '+0,50'),
        '0_72': (cfg.PLATES_0_72, 0.72, 0.48, '+0,48'),
        '0_86': (cfg.PLATES_0_86, 0.86, 0.34, '+0,34'),
    }
    if len(cfg.PLATES_0_74):
        groups_map['0_74'] = (cfg.PLATES_0_74, 0.74, 0.46, '+0,46')
    if len(cfg.PLATES_0_88):
        groups_map['0_88'] = (cfg.PLATES_0_88, 0.88, 0.32, '+0,32')
    if len(cfg.PLATES_0_48):
        groups_map['0_48'] = (cfg.PLATES_0_48, 0.48, 0.72, '+0,72')
    if len(cfg.PLATES_0_50):
        groups_map['0_50'] = (cfg.PLATES_0_50, 0.50, 0.70, '+0,70')
    if len(cfg.PLATES_0_34):
        groups_map['0_34'] = (cfg.PLATES_0_34, 0.34, 0.86, '+0,86')
    
    order = OPT_WIDTH_PRIORITY or list(groups_map.keys())
    for key in order:
        if key not in groups_map:
            continue
        items, main_w, rest_w, rest_label = groups_map[key]
        for L in items:
            sequence.append({'length': L, 'mode': 'split', 'main_w': main_w, 'rest_w': rest_w,
                             'label_main': plate_label(L, main_w), 'label_rest': rest_label})

    return sequence


# ==================== ФУНКЦИИ РИСОВАНИЯ ====================

def _draw_segment(ax, x0: float, length: float, color: str, label: str, y: float = 0.0, height: float = cfg.TRACK_WIDTH_M):
    """Рисует простой сегмент плиты"""
    rect = patches.Rectangle((x0, y), length, height, linewidth=1, edgecolor='black', facecolor=color, alpha=0.85)
    ax.add_patch(rect)
    ax.text(x0 + length/2, y + height/2, label, ha='center', va='center', fontsize=8, color='white', weight='bold')


def _draw_split_plate(ax, x0: float, length: float, main_w: float, rest_w: float, label_main: str, label_rest: str | None = None, secondary_cuts: list = None):
    """Рисует плиту с продольным резом и возможными вторичными резами в остатке"""
    # Фон всей плиты (1.2 м)
    rect = patches.Rectangle((x0, 0.0), length, cfg.TRACK_WIDTH_M, linewidth=1.2, edgecolor='black', facecolor='#ecf0f1', alpha=1.0)
    ax.add_patch(rect)
    
    # Основная часть (зелёная)
    main_rect = patches.Rectangle((x0, 0.0), length, main_w, linewidth=0.8, edgecolor='black', facecolor='#2ecc71', alpha=0.9)
    ax.add_patch(main_rect)
    
    # ПРОДОЛЬНЫЙ РЕЗ (по ширине) - СИНЯЯ ГОРИЗОНТАЛЬНАЯ ЛИНИЯ
    ax.plot([x0, x0 + length], [main_w, main_w], color='blue', linestyle='-', linewidth=2.5, alpha=0.8, zorder=10)
    
    ax.text(x0 + length/2, main_w/2, label_main, ha='center', va='center', fontsize=8, color='white', weight='bold')
    
    # Если есть вторичные резы в остатке
    if secondary_cuts and rest_w > 0.02:
        y_offset = main_w
        for i, sec_cut in enumerate(secondary_cuts):
            sec_w = sec_cut['width']
            sec_label = sec_cut['label']
            
            # ВТОРИЧНЫЙ ПРОДОЛЬНЫЙ РЕЗ - ОРАНЖЕВАЯ ЛИНИЯ (перед сегментом)
            if i == 0:
                # Первый вторичный рез (граница между первичным остатком и вторичным сегментом)
                # Линия только на длину вторичного реза
                first_sec_length = sec_cut.get('target_length', length)
                ax.plot([x0, x0 + first_sec_length], [y_offset, y_offset], color='orange', linestyle='-', linewidth=2.0, alpha=0.8, zorder=10)
            
            # Проверяем, есть ли поперечный рез для этой вторичной плиты
            if sec_cut.get('transverse_cut'):
                # Вторичная плита С поперечным резом
                target_length = sec_cut['target_length']
                remainder = sec_cut.get('remainder', 0)
                
                # Рисуем целевую часть (голубая)
                sec_rect = patches.Rectangle((x0, y_offset), target_length, sec_w, linewidth=0.8, edgecolor='black', facecolor='#3498db', alpha=0.9)
                ax.add_patch(sec_rect)
                ax.text(x0 + target_length/2, y_offset + sec_w/2, sec_label, ha='center', va='center', fontsize=7, color='white', weight='bold')
                
                # Остаток по длине (светло-серый)
                if remainder > 0.1:
                    remainder_rect = patches.Rectangle((x0 + target_length, y_offset), remainder, sec_w, 
                                                       linewidth=0.8, edgecolor='gray', facecolor='#bdc3c7', alpha=0.7)
                    ax.add_patch(remainder_rect)
                    if remainder > 0.3:
                        ax.text(x0 + target_length + remainder/2, y_offset + sec_w/2, 
                               f'ост.\n{remainder:.2f}м', ha='center', va='center', fontsize=5, color='#2c3e50')
                
                # КРАСНАЯ ВЕРТИКАЛЬНАЯ ЛИНИЯ - поперечный рез!
                ax.plot([x0 + target_length, x0 + target_length], 
                       [y_offset, y_offset + sec_w],
                       color='red', linestyle='--', linewidth=2.5, alpha=0.8, zorder=10)
            else:
                # Обычный вторичный рез (может быть с укорочением по длине)
                # Проверяем, указана ли целевая длина (для transverse-резов)
                sec_length = sec_cut.get('target_length', length)  # По умолчанию - длина основной плиты
                
                sec_rect = patches.Rectangle((x0, y_offset), sec_length, sec_w, linewidth=0.8, edgecolor='black', facecolor='#3498db', alpha=0.9)
                ax.add_patch(sec_rect)
                ax.text(x0 + sec_length/2, y_offset + sec_w/2, sec_label, ha='center', va='center', fontsize=7, color='white', weight='bold')
                
                # Если это укороченная плита, рисуем остаток по длине
                if sec_length < length - 0.1:
                    remainder_length = length - sec_length
                    remainder_rect = patches.Rectangle((x0 + sec_length, y_offset), remainder_length, sec_w,
                                                       linewidth=0.8, edgecolor='gray', facecolor='#bdc3c7', alpha=0.7)
                    ax.add_patch(remainder_rect)
                    if remainder_length > 0.3:
                        ax.text(x0 + sec_length + remainder_length/2, y_offset + sec_w/2,
                               f'ост.\n{remainder_length:.2f}м', ha='center', va='center', fontsize=5, color='#2c3e50')
                    
                    # КРАСНАЯ ВЕРТИКАЛЬНАЯ ЛИНИЯ - поперечный рез!
                    ax.plot([x0 + sec_length, x0 + sec_length],
                           [y_offset, y_offset + sec_w],
                           color='red', linestyle='--', linewidth=2.5, alpha=0.8, zorder=10)
            
            y_offset += sec_w
            
            # Линия между вторичными резами (если их несколько)
            if i < len(secondary_cuts) - 1:
                # Линия только на длину вторичного реза (не на всю основную плиту)
                sec_length = sec_cut.get('target_length', length)
                ax.plot([x0, x0 + sec_length], [y_offset, y_offset], color='orange', linestyle='-', linewidth=2.0, alpha=0.8, zorder=10)
        
        # Остаток (отход) - тёмно-серый
        remaining_w = rest_w - sum(sc['width'] for sc in secondary_cuts)
        if remaining_w > 0.01:
            # Определяем минимальную длину среди всех вторичных резов
            min_sec_length = min(sc.get('target_length', length) for sc in secondary_cuts)
            
            # Линия перед отходом (граница) - только на минимальную длину вторичных резов
            ax.plot([x0, x0 + min_sec_length], [y_offset, y_offset], color='gray', linestyle='--', linewidth=1.5, alpha=0.8, zorder=10)
            waste_rect = patches.Rectangle((x0, y_offset), min_sec_length, remaining_w, linewidth=0.5, edgecolor='gray', facecolor='#95a5a6', alpha=0.7)
            ax.add_patch(waste_rect)
            ax.text(x0 + min_sec_length/2, y_offset + remaining_w/2, f'отход {remaining_w*1000:.0f}мм', ha='center', va='center', fontsize=6, color='white')
    elif label_rest and rest_w > 0.02:
        # Обычный остаток без вторичных резов
        # Линия перед остатком уже нарисована (синяя), добавляем только метку
        ax.text(x0 + length/2, main_w + rest_w/2, label_rest, ha='center', va='center', fontsize=7, color='#2c3e50')
        # Подпись "остаток"
        ax.text(x0 + length - 0.2, main_w + rest_w/2, f'остаток\n{rest_w*1000:.0f}мм', ha='right', va='center', fontsize=6, color='#7f8c8d', style='italic')


def _draw_transverse_cut(ax, x0: float, total_length: float, target_length: float, 
                         width: float, label_target: str, remainder_length: float):
    """
    Рисует плиту с поперечным резом (по длине)
    
    ├─────────┬──────┤
    │ 3.31м   │ост.  │
    │ нужна   │0.01м │
    └─────────┴──────┘
         ↑
    поперечный рез (красная вертикальная линия)
    """
    # Фон всей плиты
    rect = patches.Rectangle((x0, 0.0), total_length, width, 
                            linewidth=1.2, edgecolor='black', 
                            facecolor='#ecf0f1', alpha=1.0)
    ax.add_patch(rect)
    
    # Левая часть (целевая плита) - зелёная
    target_rect = patches.Rectangle((x0, 0.0), target_length, width,
                                   linewidth=0.8, edgecolor='black',
                                   facecolor='#27ae60', alpha=0.9)
    ax.add_patch(target_rect)
    ax.text(x0 + target_length/2, width/2, label_target,
           ha='center', va='center', fontsize=8, color='white', weight='bold')
    
    # Правая часть (остаток) - светло-серая
    if remainder_length > 0.01:
        remainder_rect = patches.Rectangle((x0 + target_length, 0.0), 
                                          remainder_length, width,
                                          linewidth=0.8, edgecolor='gray',
                                          facecolor='#bdc3c7', alpha=0.7)
        ax.add_patch(remainder_rect)
        
        # Метка остатка по длине
        if remainder_length > 0.3:  # Показываем метку только если остаток заметный
            ax.text(x0 + target_length + remainder_length/2, width/2,
                   f'остаток\nпо длине\n{remainder_length:.2f}м',
                   ha='center', va='center', fontsize=6, color='#2c3e50', weight='bold')
    
    # КРАСНАЯ ВЕРТИКАЛЬНАЯ ЛИНИЯ - поперечный рез!
    ax.plot([x0 + target_length, x0 + target_length], 
           [0, width],
           color='red', linestyle='--', linewidth=2.5, alpha=0.8)


# ==================== ГЛАВНАЯ ФУНКЦИЯ ВИЗУАЛИЗАЦИИ ====================

def visualize_plan(output_dir: str = 'Визуализация_Раскладки'):
    """Создаёт визуализацию раскладки плит и сохраняет файлы"""
    try:
        optimized = optimize_cuts_pulp({300: 4, 500: 3, 700: 2, 900: 2})
        print("Оптимальные резы:", optimized)
    except Exception as e:
        print("[OPT] Ошибка при оптимизации:", e)

    os.makedirs(output_dir, exist_ok=True)

    price_table = load_price_table_from_xlsx(cfg.PRICE_XLSX_PATH)
    try:
        init_schema(cfg.PRICE_DB_PATH)
        written = import_from_xlsx(cfg.PRICE_XLSX_PATH, cfg.PRICE_DB_PATH)
        if written:
            print(f'[ПРАЙС->БД] записано строк: {written}')
    except Exception:
        pass

    price_rows, total_sum = build_price_rows(price_table)
    seq = build_layout_sequence()
    total_length = sum(s['length'] for s in seq)

    fig = plt.figure(figsize=(22, 14))
    gs = fig.add_gridspec(4, 1, height_ratios=[3.0, 1.0, 1.4, 1.8])
    ax_track = fig.add_subplot(gs[0, 0])
    ax_strips = fig.add_subplot(gs[1, 0])
    ax_table = fig.add_subplot(gs[2, 0])
    ax_price = fig.add_subplot(gs[3, 0])
    fig.suptitle('КЗ: Дорожка 1 (ширина 1.2 м) — раскладка, резы, ведомости и смета', fontsize=16, fontweight='bold')

    ax_track.set_xlim(0, max(total_length + 2, cfg.TRACK_LENGTH_M))
    ax_track.set_ylim(0, cfg.TRACK_WIDTH_M + 0.8)
    ax_track.set_aspect('auto')
    ax_track.spines['top'].set_visible(False)
    ax_track.spines['right'].set_visible(False)
    ax_track.set_yticks([0, 0.6, 1.2])
    ax_track.set_yticklabels(['0', '0.6', '1.2 м'])
    ax_track.set_xlabel('Длина (м)')
    ax_track.set_xticks(range(0, int(max(total_length, cfg.TRACK_LENGTH_M)) + 1, 5))
    ax_track.grid(axis='x', linestyle=':', linewidth=0.5, alpha=0.5)

    track_rect = patches.Rectangle((0, 0), cfg.TRACK_LENGTH_M, cfg.TRACK_WIDTH_M, linewidth=2, edgecolor='black', facecolor='none', linestyle='--')
    ax_track.add_patch(track_rect)

    x = 0.0
    for item in seq:
        if item.get('mode') == 'solid':
            _draw_segment(ax_track, x, item['length'], '#2ecc71', item['label'])
        elif item.get('mode') == 'transverse':
            # Плита с поперечным резом (по длине)
            _draw_transverse_cut(
                ax_track, x, 
                total_length=item['length'],
                target_length=item['target_length'],
                width=item['width'],
                label_target=item['label_target'],
                remainder_length=item['remainder']
            )
        else:
            # Плиты с резами (первичными и возможными вторичными)
            _draw_split_plate(
                ax_track, x, item['length'],
                main_w=item['main_w'], rest_w=item['rest_w'],
                label_main=item['label_main'], label_rest=item.get('label_rest'),
                secondary_cuts=item.get('secondary_cuts')
            )
        x += item['length']

    legend_patches = [
        patches.Patch(facecolor='#2ecc71', edgecolor='black', label='🟢 Основа (первичный рез)'),
        patches.Patch(facecolor='#3498db', edgecolor='black', label='🔵 Вторичный рез (из остатка)'),
        patches.Patch(facecolor='#95a5a6', edgecolor='gray', label='⬛ Отход'),
        patches.Patch(facecolor='#ecf0f1', edgecolor='black', label='⬜ Остаток (не использован)'),
        Line2D([0], [0], color='blue', linestyle='-', linewidth=2.5, label='━ Продольный рез (первичный)'),
        Line2D([0], [0], color='orange', linestyle='-', linewidth=2.0, label='━ Продольный рез (вторичный)'),
        Line2D([0], [0], color='red', linestyle='--', linewidth=2.5, label='┊ Поперечный рез (по длине)'),
    ]
    ax_track.legend(handles=legend_patches, loc='upper right', fontsize=9)

    ax_strips.set_xlim(0, 100)
    ax_strips.set_ylim(0, 1)
    ax_strips.axis('off')

    # Формируем сводку с учётом каскадной оптимизации
    from optimization import OPT_CASCADING_PLAN
    txt = (
        f"Длина по плану: {total_length:.1f} м  |  Продольных резов: {cfg.LONGITUDINAL_CUTS}  |  Подрезов по длине: {cfg.LENGTH_TRIMS}\n"
        f"Остатки лент 0.3: {cfg.UNUSED_STRIPS_0_3_M_TOTAL:.1f} пог.м  |  Обрезки 0.2: {cfg.SCRAP_STRIPS_0_2_M_TOTAL:.1f} пог.м (≈ {cfg.WASTE_AREA_M2:.2f} м²)"
    )
    
    # Добавляем информацию о каскадной оптимизации, если она была использована
    if OPT_CASCADING_PLAN and OPT_CASCADING_PLAN.get('total_plates', 0) > 0:
        txt += f"\n\nОПТИМИЗАЦИЯ: Плит потребуется {OPT_CASCADING_PLAN['total_plates']} шт (с каскадными резами)"
        txt += f" | Отходы: {OPT_CASCADING_PLAN.get('waste_width', 0)} мм"
    
    ax_strips.text(0.02, 0.6, txt, ha='left', va='center', fontsize=12,
                   bbox=dict(boxstyle='round,pad=0.6', facecolor='#f8f9fa', edgecolor='#bdc3c7'))

    # Формируем детальный план резов или остатков
    if OPT_CASCADING_PLAN and OPT_CASCADING_PLAN.get('total_plates', 0) > 0:
        # Показываем детальный план каскадных резов
        details = "[PLAN] ДЕТАЛЬНЫЙ ПЛАН РЕЗОВ:\n\n"
        
        # Первичные резы
        if OPT_CASCADING_PLAN.get('primary_cuts'):
            details += "[1] Первичные резы (из 1200 мм):\n"
            for cut in OPT_CASCADING_PLAN['primary_cuts']:
                details += f"  • {cut['qty']} плит → {cut['width']} мм + остаток {cut['rest']} мм\n"
        
        # Вторичные резы
        if OPT_CASCADING_PLAN.get('secondary_cuts'):
            details += "\n[2] Вторичные резы (из остатков):\n"
            for cut in OPT_CASCADING_PLAN['secondary_cuts']:
                if cut.get('pieces', 1) > 1:
                    details += f"  • {cut['qty']} остатков {cut['source']} мм → {cut['pieces']} частей по {cut['cuts'][0]} мм"
                    if cut.get('waste', 0) > 0:
                        details += f" (отход {cut['waste']} мм)"
                    details += "\n"
                else:
                    cuts_str = ' + '.join(str(c) for c in cut['cuts'])
                    details += f"  • {cut['qty']} остатков {cut['source']} мм → {cuts_str} мм"
                    if cut.get('waste', 0) > 0:
                        details += f" (отход {cut['waste']} мм)"
                    details += "\n"
        
        # Поперечные резы
        if OPT_CASCADING_PLAN.get('transverse_cuts'):
            details += "\n[RED] Поперечные резы (по длине):\n"
            for tcut in OPT_CASCADING_PLAN['transverse_cuts']:
                details += f"  • Плита {tcut['source_length']}м x {tcut['source_width']}мм -> {tcut['target_length']}м"
                if tcut.get('remainder', 0) > 0.1:
                    details += f" (остаток {tcut['remainder']:.2f}м)"
                details += "\n"
        
        ax_strips.text(0.02, 0.15, details, ha='left', va='center', fontsize=10,
                       bbox=dict(boxstyle='round,pad=0.5', facecolor='#e8f5e9', edgecolor='#66bb6a'))
    else:
        # Fallback: показываем старую информацию об остатках
        leftovers = (
            "Остатки/обрезки:\n"
            "Ленты 0.3: 3x3.8 м; 1x2.9 м\n"
            f"Ленты 0.2 (обрезки): {', '.join(f'{L:.1f} м' for L in cfg.PLATES_1_0)}"
        )
        ax_strips.text(0.02, 0.15, leftovers, ha='left', va='center', fontsize=11,
                       bbox=dict(boxstyle='round,pad=0.5', facecolor='#eef7ff', edgecolor='#a3c9ff'))

    ax_table.axis('off')

    # Формируем список заказа из реальных данных
    order_list = []
    from collections import Counter
    
    # Собираем все плиты из заказа
    all_orders = []
    for width_mm, plates_list in [
        (320, cfg.PLATES_0_32), (460, cfg.PLATES_0_46), (700, cfg.PLATES_0_70),
        (720, cfg.PLATES_0_72), (860, cfg.PLATES_0_86), (880, cfg.PLATES_0_88),
        (740, cfg.PLATES_0_74), (480, cfg.PLATES_0_48), (500, cfg.PLATES_0_50),
        (340, cfg.PLATES_0_34), (1080, cfg.PLATES_1_08)
    ]:
        if plates_list:
            length_counts = Counter(plates_list)
            for length, qty in sorted(length_counts.items(), key=lambda x: (-x[0], -x[1])):
                all_orders.append({
                    'length': length,
                    'width': width_mm,
                    'qty': qty
                })
    
    # Формируем строки заказа
    if all_orders:
        # Группируем по длинам для компактности
        for order in all_orders:
            order_list.append(f"Заказ {order['length']:.1f}мx{order['width']}мм: {order['qty']} шт")
    else:
        order_list.append('Заказ не найден')
    
    # Формируем список использования из оптимизации
    used_list = []
    
    if OPT_CASCADING_PLAN and OPT_CASCADING_PLAN.get('total_plates', 0) > 0:
        # Итого плит
        used_list.append(f"Плит 1200мм потребуется: {OPT_CASCADING_PLAN['total_plates']} шт")
        
        # Первичные резы
        if OPT_CASCADING_PLAN.get('primary_cuts'):
            primary_info = []
            for cut in OPT_CASCADING_PLAN['primary_cuts']:
                primary_info.append(f"{cut['qty']}x({cut['width']}мм+{cut['rest']}мм)")
            if primary_info:
                used_list.append(f"Первичные резы: {'; '.join(primary_info)}")
        
        # Вторичные резы
        if OPT_CASCADING_PLAN.get('secondary_cuts'):
            secondary_info = []
            for cut in OPT_CASCADING_PLAN['secondary_cuts']:
                if cut.get('pieces', 1) > 1:
                    secondary_info.append(f"{cut['qty']}x{cut['source']}мм->{cut['pieces']}x{cut['cuts'][0]}мм")
                else:
                    secondary_info.append(f"{cut['qty']}x{cut['source']}мм->{cut['cuts'][0]}мм")
            if secondary_info:
                used_list.append(f"Вторичные резы: {'; '.join(secondary_info)}")
        
        # Поперечные резы
        if OPT_CASCADING_PLAN.get('transverse_cuts'):
            trans_count = len(OPT_CASCADING_PLAN['transverse_cuts'])
            used_list.append(f"Поперечных резов: {trans_count}")
        
        # Отходы
        if OPT_CASCADING_PLAN.get('waste_width', 0) > 0:
            used_list.append(f"Отходы: {OPT_CASCADING_PLAN['waste_width']} мм по ширине")
    else:
        # Старый формат, если оптимизация не использовалась
        used_list.append('1.2 без реза: 6.3x2; 3.8x2')
        used_list.append('1.5->1.2: 3.8x3; 2.9x1 (остаток 0.3)')
        used_list.append(f'Резы: продольных {cfg.LONGITUDINAL_CUTS}; подрезов {cfg.LENGTH_TRIMS}')

    rows = max(len(order_list), len(used_list))
    table_rows = []
    for i in range(rows):
        left = order_list[i] if i < len(order_list) else ''
        right = used_list[i] if i < len(used_list) else ''
        table_rows.append([left, right])

    col_labels = ['Список плит по заказу', 'Использовано (с учётом резов) / остатки / обрезки']

    table = ax_table.table(cellText=table_rows, colLabels=col_labels, loc='center', cellLoc='left', colLoc='center')
    table.auto_set_font_size(False)
    table.set_fontsize(11)
    table.scale(1, 1.5)

    ax_price.axis('off')
    price_headers = ['№', 'Наименование', 'Кол-во', 'Ед.', 'Неделя', 'Контрагент', 'Вес(кг)', 'Цена', 'Сумма']
    price_table = ax_price.table(cellText=price_rows, colLabels=price_headers, loc='center', cellLoc='center', colLoc='center')
    price_table.auto_set_font_size(False)
    price_table.set_fontsize(10)
    price_table.scale(1, 1.4)
    price_col_idx = price_headers.index('Цена')
    not_priced = any(row[price_col_idx].strip().startswith('0') for row in price_rows)
    
    # Используем стоимость из каскадной оптимизации, если она доступна
    if OPT_CASCADING_PLAN and OPT_CASCADING_PLAN.get('total_cost', 0) > 0:
        optimized_cost = OPT_CASCADING_PLAN['total_cost']
        title = f'Итоговая стоимость: {optimized_cost:,.2f} ₽ (оптимизировано)'.replace(',', ' ').replace('.', ',')
    else:
        title = f'Итоговая стоимость: {total_sum:,.2f} ₽'.replace(',', ' ').replace('.', ',')
    
    if not_priced:
        title += ' (внимание: не найдены цены для некоторых позиций — проверьте прайс)'
    ax_price.set_title(title, fontsize=12, pad=10)

    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S_%f')
    csv_path = os.path.join(output_dir, f'Ведомость_Дорожка_1_{timestamp}.csv')
    with open(csv_path, 'w', encoding='utf-8') as f:
        f.write('Список плит по заказу;Использовано (с учётом резов) / остатки / обрезки\n')
        for left, right in table_rows:
            f.write(f'{left};{right}\n')

    if pd is not None:
        try:
            df_v = pd.DataFrame(table_rows, columns=col_labels)
            xlsx_path_v = os.path.join(output_dir, f'Ведомость_Дорожка_1_{timestamp}.xlsx')
            df_v.to_excel(xlsx_path_v, index=False)

            df_p = pd.DataFrame(price_rows, columns=price_headers)
            xlsx_path_p = os.path.join(output_dir, f'Смета_Дорожка_1_{timestamp}.xlsx')
            with pd.ExcelWriter(xlsx_path_p, engine='openpyxl') as writer:
                df_p.to_excel(writer, index=False, sheet_name='Смета')
                df_v.to_excel(writer, index=False, sheet_name='Ведомость')
        except Exception:
            pass

    png_path = os.path.join(output_dir, f'Схема_Дорожка_1_КЗ_{timestamp}.png')
    pdf_path = os.path.join(output_dir, f'Схема_Дорожка_1_КЗ_{timestamp}.pdf')
    plt.tight_layout(rect=[0, 0, 1, 0.95])
    fig.savefig(png_path, dpi=300)
    fig.savefig(pdf_path)
    plt.close(fig)

    print('[ГОТОВО] Визуализация и файлы сохранены:')
    print('  PNG:', png_path)
    print('  PDF:', pdf_path)
    print('  CSV:', csv_path)
    if pd is not None:
        print('  XLSX (ведомость):', os.path.join(output_dir, f'Ведомость_Дорожка_1_{timestamp}.xlsx'))
        print('  XLSX (смета):', os.path.join(output_dir, f'Смета_Дорожка_1_{timestamp}.xlsx'))
    return png_path, pdf_path


if __name__ == '__main__':
    visualize_plan()



