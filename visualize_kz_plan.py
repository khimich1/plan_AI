#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Визуализация КЗ-плана для одной дорожки 1.2 м:
- Укладка сегментов 1.2
- Преобразование 1.5 -> 1.2 + 0.3
- Получение 1.0 из 1.2 резом: 1.2 -> 1.0 + 0.2 (0.2 в обрезки)
- Показ суммарных резов и остатков/обрезков
- Расчёт стоимости: берём цены плит из XLSX `банк знаний/Новые цены для прайса с 19.08.24.xlsx`
  и цену реза из DOCX `банк знаний/Письмо Цены с 29.05.2024 цены на резы.docx` (если доступно)
Результат: PNG и PDF в папке "Визуализация_Раскладки".
Также выгружаются CSV/XLSX с ведомостью и сметой.
"""
import os
from datetime import datetime
import re
import math
import matplotlib
matplotlib.use('Agg')  # headless backend to avoid GUI/thread issues
import matplotlib.pyplot as plt
import matplotlib.patches as patches
import sqlite3
from price_db import init_schema, import_from_xlsx, get_price

try:
    import pandas as pd  # для XLSX (опционально)
except Exception:
    pd = None

try:
    from docx import Document  # чтение цены резов из DOCX (опционально)
except Exception:
    Document = None

TRACK_LENGTH_M = 101.0
TRACK_WIDTH_M = 1.2

# Пути к прайсам (делаем абсолютными относительно файла скрипта)
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
PRICE_XLSX_PATH = os.path.join(BASE_DIR, 'банк знаний', 'Новые цены для прайса с 19.08.24.xlsx')
CUTS_DOCX_PATH = os.path.join(BASE_DIR, 'банк знаний', 'Письмо Цены с 29.05.2024 цены на резы.docx')  # не используется в новой модели
PRICE_DB_PATH = os.path.join(BASE_DIR, 'pb.db')

# Стоимость резов
LONG_CUT_PRICE_PER_M = 460.0  # Продольный рез, руб/пог.м
TRANSVERSE_CUT_PRICE = 1200.0  # Поперечный (или скошенный) рез, руб/шт

# Данные из согласованного КЗ-плана
# 1) Плиты 1.2 м — без резов (новый заказ)
#  - ПБ 37,9-4,6-8п ×4 ×3 строки = 12 шт  → длина 3.79 м (ТРЕБУЕТ рез до 0.46, см. ниже)
#  - ПБ 58,6-10,8-8п ×5         = 5 шт   → длина 5.86 м (ТРЕБУЕТ рез до 1.08, см. ниже)
#  - ПБ 33,9-7,2-8п ×2         = 2 шт   → длина 3.39 м
# Для визуализации «как есть» оставляем только 3.39 как 1.2 м.
PLATES_1_2 = [3.39]*2

# Дополнительные целевые ширины, которые получаем продольным резом из 1.2 м
#  - 1.08 м: диапазон 1.02–1.08 (остаток 0.12–0.18)
#  - 0.46 м: диапазон 0.46–0.53 (остаток 0.67–0.74)
#  - 0.32 м: диапазон 0.26–0.32 (остаток 0.88–0.94)
#  - 0.72 м: диапазон 0.66–0.72 (остаток 0.48–0.54)
#  - 0.70 м: в рамках 0.66–0.72 (остаток ~0.50)
#  - 0.86 м: диапазон 0.86–0.92 (остаток 0.28–0.34)
PLATES_1_08 = []         # нет 1.08 в этом заказе
PLATES_0_46 = []         # нет 0.46 в этом заказе
PLATES_0_32 = [6.63]*4 + [7.83]*3
PLATES_0_72 = [5.63]*5
PLATES_0_70 = [4.65]*5
PLATES_0_86 = [6.75]*2 + [4.65]*5
# Заказы на вторую половину (если пользователь прислал такие ширины)
PLATES_0_74 = []
PLATES_0_88 = []
PLATES_0_48 = []
PLATES_0_50 = []
PLATES_0_34 = []
# 2) Плиты 1.5 м — используем как 1.2 м (лента 0.3 образуется)
PLATES_1_5_TO_1_2 = []
# 3) Плиты 1.0 м — получаем из 1.2 (остаток 0.2 уходит в обрезки)
PLATES_1_0 = []

# Резы по плану: по одному на каждую плиту, получаемую резом
LONGITUDINAL_CUTS = (
    len(PLATES_1_5_TO_1_2) + len(PLATES_1_0) +
    len(PLATES_1_08) + len(PLATES_0_46) +
    len(PLATES_0_32) + len(PLATES_0_72) + len(PLATES_0_70) + len(PLATES_0_86)
)
LENGTH_TRIMS = 0

# Остатки и отходы
UNUSED_STRIPS_0_3_M_TOTAL = 0.0
SCRAP_STRIPS_0_2_M_TOTAL = 0.0
# Новые ленты/обрезки от резов 1.2 -> 1.08 (0.12 в обрезки) и 1.2 -> 0.46 (0.74 как используемая лента)
USABLE_STRIPS_0_74_M_TOTAL = round(sum(PLATES_0_46), 1)
USABLE_STRIPS_0_88_M_TOTAL = round(sum(PLATES_0_32), 1)
USABLE_STRIPS_0_48_M_TOTAL = round(sum(PLATES_0_72), 1)
USABLE_STRIPS_0_50_M_TOTAL = round(sum(PLATES_0_70), 1)
USABLE_STRIPS_0_34_M_TOTAL = round(sum(PLATES_0_86), 1)
SCRAP_STRIPS_0_12_M_TOTAL = round(sum(PLATES_1_08), 1)
WASTE_AREA_M2 = round(0.12 * SCRAP_STRIPS_0_12_M_TOTAL, 2)


# ---------- Утилиты для прайса ----------

# ---------- Парсинг и настройка списков плит из текста пользователя ----------

def _clear_all_plate_lists():
    global PLATES_1_2, PLATES_1_5_TO_1_2, PLATES_1_0, PLATES_1_08
    global PLATES_0_46, PLATES_0_32, PLATES_0_72, PLATES_0_70, PLATES_0_86
    global PLATES_0_74, PLATES_0_88, PLATES_0_48, PLATES_0_50, PLATES_0_34
    PLATES_1_2 = []
    PLATES_1_5_TO_1_2 = []
    PLATES_1_0 = []
    PLATES_1_08 = []
    PLATES_0_46 = []
    PLATES_0_32 = []
    PLATES_0_72 = []
    PLATES_0_70 = []
    PLATES_0_86 = []
    PLATES_0_74 = []
    PLATES_0_88 = []
    PLATES_0_48 = []
    PLATES_0_50 = []
    PLATES_0_34 = []


def _recompute_totals_from_lists():
    global LONGITUDINAL_CUTS, LENGTH_TRIMS
    global UNUSED_STRIPS_0_3_M_TOTAL, SCRAP_STRIPS_0_2_M_TOTAL
    global USABLE_STRIPS_0_74_M_TOTAL, USABLE_STRIPS_0_88_M_TOTAL
    global USABLE_STRIPS_0_48_M_TOTAL, USABLE_STRIPS_0_50_M_TOTAL
    global USABLE_STRIPS_0_34_M_TOTAL, SCRAP_STRIPS_0_12_M_TOTAL
    global WASTE_AREA_M2

    LONGITUDINAL_CUTS = (
        len(PLATES_1_5_TO_1_2) + len(PLATES_1_0) +
        len(PLATES_1_08) + len(PLATES_0_46) +
        len(PLATES_0_32) + len(PLATES_0_72) + len(PLATES_0_70) + len(PLATES_0_86)
    )
    LENGTH_TRIMS = 0

    UNUSED_STRIPS_0_3_M_TOTAL = 0.0
    SCRAP_STRIPS_0_2_M_TOTAL = 0.0
    USABLE_STRIPS_0_74_M_TOTAL = round(sum(PLATES_0_46), 1)
    USABLE_STRIPS_0_88_M_TOTAL = round(sum(PLATES_0_32), 1)
    USABLE_STRIPS_0_48_M_TOTAL = round(sum(PLATES_0_72), 1)
    USABLE_STRIPS_0_50_M_TOTAL = round(sum(PLATES_0_70), 1)
    USABLE_STRIPS_0_34_M_TOTAL = round(sum(PLATES_0_86), 1)
    SCRAP_STRIPS_0_12_M_TOTAL = round(sum(PLATES_1_08), 1)
    WASTE_AREA_M2 = round(0.12 * SCRAP_STRIPS_0_12_M_TOTAL, 2)


def set_plate_lists_from_text(user_text: str) -> None:
    """Парсит свободный текст пользователя и заполняет списки PLATES_*.

    Поддерживаем форматы:
      - "1.2×3.39 — 2 шт" / "0,32x6,63 - 4"
      - "Плиты ПБ 78-12-8п 3" (длина в дм, ширина 12 => 1.2м, количество 3)
    Неизвестные ширины игнорируем.
    """
    _clear_all_plate_lists()

    text = (user_text or '').replace('\u00d7', 'x').replace('×', 'x')
    lines = [l.strip() for l in re.split(r'[\n;]+', text) if l.strip()]

    def add_items(width_m: float, length_m: float, qty: int):
        # Специальная обработка плит 1.5 м → заменяем на 1.2 м + 0.3 м
        if 1.45 <= width_m <= 1.55:  # 1.5 м (диапазон ±50 мм)
            # Добавляем плиту 1.2 м
            for _ in range(max(0, qty)):
                PLATES_1_2.append(round(float(length_m), 2))
            # Добавляем плиту 0.3 м (записываем в PLATES_0_32)
            for _ in range(max(0, qty)):
                PLATES_0_32.append(round(float(length_m), 2))
            return
        
        target = None
        if 1.15 <= width_m <= 1.25:
            target = PLATES_1_2
        elif 0.98 <= width_m <= 1.02:
            target = PLATES_1_0
        elif 1.06 <= width_m <= 1.12:
            target = PLATES_1_08
        elif 0.30 <= width_m <= 0.33:
            target = PLATES_0_32
        elif 0.44 <= width_m <= 0.47:
            target = PLATES_0_46
        elif 0.69 <= width_m <= 0.71:
            target = PLATES_0_70
        elif 0.71 < width_m <= 0.73:
            target = PLATES_0_72
        elif 0.85 <= width_m <= 0.865:
            target = PLATES_0_86
        # также позволяем напрямую заказывать вторые половины
        elif 0.33 < width_m <= 0.35:
            target = PLATES_0_34
        elif 0.47 < width_m <= 0.49:
            target = PLATES_0_48
        elif 0.49 < width_m <= 0.51:
            target = PLATES_0_50
        elif 0.73 < width_m <= 0.75:
            target = PLATES_0_74
        elif 0.865 < width_m <= 0.895:
            target = PLATES_0_88
        else:
            return
        for _ in range(max(0, qty)):
            target.append(round(float(length_m), 2))

    for raw in lines:
        s = raw.lower()
        # 1) формат WxL x qty (поддерживает запятую и точку)
        s_norm = s.replace(',', '.')
        m = re.search(r'(\d+(?:\.\d+)?)\s*[xх]\s*(\d+(?:\.\d+)?)\D*(\d+)?', s)
        if m:
            w = float(m.group(1).replace(',', '.'))
            L = float(m.group(2).replace(',', '.'))
            q = int((m.group(3) or '1').replace(',', '.'))
            add_items(w, L, q)
            continue
        # 2) формат "Плиты ПБ 78,3-3,2-8п 3" или "ПБ 78-12-8п 10"
        m2 = re.search(r'плиты?\s*пб\s*([\d\.,]+)\s*-\s*([\d\.,]+)', s)
        if not m2:
            m2 = re.search(r'\bпб\s*([\d\.,]+)\s*-\s*([\d\.,]+)', s)
        if m2:
            Ldm_str = m2.group(1).replace(' ', '').replace(',', '.')
            Wdm_str = m2.group(2).replace(' ', '').replace(',', '.')
            try:
                # В нотации ПБ длина указывается в дм. Всегда делим на 10 → метры
                L = float(Ldm_str) / 10.0
            except Exception:
                continue
            try:
                # Ширина также в дм (12 → 1.2; 3.2 → 0.32; 8.6 → 0.86)
                W = float(Wdm_str) / 10.0
            except Exception:
                continue
            q = 1
            # Количество — последнее число в строке
            mq = re.search(r'(\d+)\s*(шт)?\s*$', s)
            if mq:
                try:
                    q = int(mq.group(1))
                except Exception:
                    q = 1
            add_items(W, L, q)
            continue

    _recompute_totals_from_lists()


# ---------- Использование оптимизатора для порядка ширин ----------

OPT_WIDTH_PRIORITY: list[str] = []  # например: ['0_32','0_46','0_70','0_72','0_86']
OPT_PLAN: dict = {}  # результат полной оптимизации: как закрывать спрос

def optimize_full_plan_with_narrowing() -> dict:
    """Полная ILP-оптимизация с учётом источников: split 1.2, narrowing (сужение), trans cut.
    Минимизирует стоимость = цена плит + резы + отходы + штраф за неиспользованные остатки.
    Возвращает plan: {'actions': [(source_type, W_main, W_pair/src, L, qty, long_cuts, trans_cuts), ...]}
    """
    try:
        from pulp import LpProblem, LpMinimize, LpVariable, LpInteger, lpSum, value
    except Exception:
        print('[OPT_FULL] PuLP не установлен, пропускаем.')
        return {}
    
    # 1) Спрос: {(W_mm, L_m): qty}
    demand = {}
    def add_demand(w_mm: int, lengths: list[float]):
        for L in lengths:
            key = (w_mm, round(L, 2))
            demand[key] = demand.get(key, 0) + 1
    add_demand(1200, PLATES_1_2)
    add_demand(320, PLATES_0_32); add_demand(460, PLATES_0_46)
    add_demand(720, PLATES_0_72); add_demand(700, PLATES_0_70); add_demand(860, PLATES_0_86)
    add_demand(880, PLATES_0_88); add_demand(740, PLATES_0_74)
    add_demand(480, PLATES_0_48); add_demand(500, PLATES_0_50); add_demand(340, PLATES_0_34)
    
    if not demand:
        return {}
    
    # 2) Источники
    split_pairs = [(320,880),(460,740),(720,480),(700,500),(860,340)]
    narrowing_options = [(340,320,20),(500,480,20),(740,720,20),(880,860,20),(480,460,20)]
    lengths_set = sorted(set(L for (W,L) in demand.keys()))
    
    prob = LpProblem('full_narrow', LpMinimize)
    x_split = {}; x_narrow = {}; x_solid = {}
    
    for (Wm, Wp) in split_pairs:
        for L in lengths_set:
            x_split[(Wm, Wp, L)] = LpVariable(f"sp_{Wm}_{Wp}_{L}", lowBound=0, cat=LpInteger)
    for (Wsrc, Wtgt, delta) in narrowing_options:
        for L in lengths_set:
            x_narrow[(Wsrc, Wtgt, L)] = LpVariable(f"nr_{Wsrc}_{Wtgt}_{L}", lowBound=0, cat=LpInteger)
    for L in lengths_set:
        x_solid[L] = LpVariable(f"sol_1200_{L}", lowBound=0, cat=LpInteger)
    
    # Спрос покрыт
    for (W, L), qty in demand.items():
        sources = []
        # split даёт основную ширину W
        for (Wm, Wp) in split_pairs:
            if Wm == W:
                sources.append(x_split.get((Wm, Wp, L), 0))
        # narrowing даёт целевую W
        for (Wsrc, Wtgt, delta) in narrowing_options:
            if Wtgt == W:
                sources.append(x_narrow.get((Wsrc, W, L), 0))
        # solid для 1.2
        if W == 1200:
            sources.append(x_solid.get(L, 0))
        # Также: если W — это парная ширина Wp от split, она тоже может покрыть спрос
        for (Wm, Wp) in split_pairs:
            if Wp == W:
                sources.append(x_split.get((Wm, Wp, L), 0))
        if sources:
            prob += lpSum(sources) >= qty, f"d_{W}_{L}"
    
    # Баланс: для narrowing нужен источник Wsrc
    # Источник Wsrc×L = остаток от split (Wp=Wsrc) минус прямое использование Wsrc
    for (Wsrc, Wtgt, delta) in narrowing_options:
        for L in lengths_set:
            needed = x_narrow.get((Wsrc, Wtgt, L), 0)
            # Источники Wsrc×L:
            produced = []
            # split даёт Wsrc как парный (Wp)
            for (Wm, Wp) in split_pairs:
                if Wp == Wsrc:
                    produced.append(x_split.get((Wm, Wp, L), 0))
            # Wsrc может покрывать прямой спрос (если он есть)
            direct_use = 0
            if (Wsrc, L) in demand:
                # Спрос на Wsrc покрывается split-pair (Wm,Wsrc)
                for (Wm, Wp) in split_pairs:
                    if Wp == Wsrc:
                        direct_use = demand[(Wsrc, L)]
                        break
            # Остаток = produced - direct_use
            # Баланс: produced - direct_use >= needed
            if produced:
                prob += lpSum(produced) >= needed + direct_use, f"src_{Wsrc}_{Wtgt}_{L}"
    
    # Стоимость
    cost = 0
    for (Wm, Wp) in split_pairs:
        for L in lengths_set:
            xvar = x_split.get((Wm, Wp, L), 0)
            plate_price = get_price(L, 8, PRICE_DB_PATH) or 10000
            cut_cost = LONG_CUT_PRICE_PER_M * L
            cost += xvar * (plate_price + cut_cost)
    for (Wsrc, Wtgt, delta) in narrowing_options:
        for L in lengths_set:
            xvar = x_narrow.get((Wsrc, Wtgt, L), 0)
            cut_cost = LONG_CUT_PRICE_PER_M * L
            waste_cost = (delta / 1200.0) * (get_price(L, 6, PRICE_DB_PATH) or 5000)
            cost += xvar * (cut_cost + waste_cost)
    for L in lengths_set:
        xvar = x_solid.get(L, 0)
        plate_price = get_price(L, 8, PRICE_DB_PATH) or 10000
        cost += xvar * plate_price
    
    # Штраф за неприспособленные остатки: считаем остатки Wp, которые не пошли дальше
    for (Wm, Wp) in split_pairs:
        for L in lengths_set:
            produced = x_split.get((Wm, Wp, L), 0)
            used = 0
            # Wp использован в split
            for (Wm2, Wp2) in split_pairs:
                if Wm2 == Wp:
                    used += x_split.get((Wm2, Wp2, L), 0)
            # Wp использован в narrowing
            for (Wsrc, Wtgt, delta) in narrowing_options:
                if Wsrc == Wp:
                    used += x_narrow.get((Wsrc, Wtgt, L), 0)
            # Wp покрыл прямой спрос (если есть)
            if (Wp, L) in demand:
                used += 0  # уже учтено в спросе
            # Остаток: produced - used
            unused_var = LpVariable(f"unused_{Wp}_{L}", lowBound=0, cat=LpInteger)
            prob += unused_var >= produced - used, f"unused_bal_{Wp}_{L}"
            cost += 3000 * unused_var
    
    prob += cost
    prob.solve()
    
    actions = []
    for (Wm, Wp) in split_pairs:
        for L in lengths_set:
            try:
                qty = int(round(value(x_split.get((Wm, Wp, L), 0))))
                if qty > 0:
                    actions.append(('split', Wm, Wp, L, qty, 1, 0))
            except: pass
    for (Wsrc, Wtgt, delta) in narrowing_options:
        for L in lengths_set:
            try:
                qty = int(round(value(x_narrow.get((Wsrc, Wtgt, L), 0))))
                if qty > 0:
                    actions.append(('narrow', Wtgt, Wsrc, L, qty, 1, 0))
            except: pass
    for L in lengths_set:
        try:
            qty = int(round(value(x_solid.get(L, 0))))
            if qty > 0:
                actions.append(('solid', 1200, 0, L, qty, 0, 0))
        except: pass
    
    try:
        total_cost = value(cost)
    except:
        total_cost = 0
    return {'actions': actions, 'summary': {'total_cost': total_cost}}


def apply_width_optimization() -> None:
    """Формирует приоритет ширин на основе mini-optimizer по спросу.
    Мы агрегируем только по ширине (<1.2) и используем результат, чтобы
    упорядочить вывод в визуализации. Длины сохраняем исходные.
    """
    global OPT_WIDTH_PRIORITY, OPT_PLAN
    
    # Сначала полная оптимизация с narrowing
    OPT_PLAN = optimize_full_plan_with_narrowing()
    if OPT_PLAN and OPT_PLAN.get('actions'):
        print(f"[OPT_FULL] Найдено {len(OPT_PLAN['actions'])} действий, стоимость: {OPT_PLAN['summary']['total_cost']:.2f}")
        # Логируем примеры
        for act in OPT_PLAN['actions'][:5]:
            src_type, W1, W2, L, qty, lc, tc = act
            print(f"  {src_type}: W={W1}mm (парный={W2}mm), L={L}м, qty={qty}, резы: long={lc}, trans={tc}")
    else:
        print("[OPT_FULL] Оптимизация не дала результата, продолжаем базовым путём.")
    orders = {}
    def add(mm: int, n: int):
        if n > 0:
            orders[mm] = orders.get(mm, 0) + n
    add(300, len(PLATES_0_32))
    add(500, len(PLATES_0_46))
    # 700 группа покрывает 0.70 и 0.72, но также учитываем спрос на их пары
    add(700, len(PLATES_0_70) + len(PLATES_0_72) + len(PLATES_0_50) + len(PLATES_0_48))
    add(900, len(PLATES_0_86) + len(PLATES_0_34))
    pulp_result = optimize_cuts_pulp(orders)
    # Базовый приоритет от оптимизатора по ширине
    base = []
    for r in pulp_result:
        if (r.get('qty') or 0) <= 0:
            continue
        cid = r.get('cut_id')
        if cid == 'cut300':
            base.append('0_32')
        elif cid == 'cut500':
            base.append('0_46')
        elif cid == 'cut700':
            base.extend(['0_70', '0_72'])
        elif cid == 'cut900':
            base.append('0_86')
    for k in ['0_32','0_46','0_70','0_72','0_86']:
        if k not in base:
            base.append(k)

    # Усиливаем приоритет с учётом совпадений длин (меньше поперечных резов лучше)
    def match_ratio(a: list[float], b: list[float]) -> float:
        if not a or not b:
            return 0.0
        asorted = sorted(round(x, 2) for x in a)
        bsorted = sorted(round(x, 2) for x in b)
        i = j = matches = 0
        while i < len(asorted) and j < len(bsorted):
            if abs(asorted[i] - bsorted[j]) <= 0.05:
                matches += 1; i += 1; j += 1
            elif asorted[i] < bsorted[j]:
                i += 1
            else:
                j += 1
        return matches / min(len(a), len(b)) if min(len(a), len(b)) else 0.0

    # ILP по длинам: собираем вход
    width_demand = {
        320: list(PLATES_0_32),
        460: list(PLATES_0_46),
        700: list(PLATES_0_70),
        720: list(PLATES_0_72),
        860: list(PLATES_0_86),
        880: list(PLATES_0_88),
        740: list(PLATES_0_74),
        480: list(PLATES_0_48),
        500: list(PLATES_0_50),
        340: list(PLATES_0_34),
    }
    pair_map = {320:880, 460:740, 700:500, 720:480, 860:340}
    len_result = optimize_with_lengths(width_demand, pair_map, trans_cut_penalty=0.6)
    # Преобразуем результат в score: выше — лучше (меньше поперечных)
    def score_from_len(w_key: str) -> float:
        mm = {'0_32':320,'0_46':460,'0_70':700,'0_72':720,'0_86':860}[w_key]
        r = len_result.get(mm, {})
        matched = float(r.get('matched', 0))
        trans = float(r.get('trans_cuts', 0))
        total = float(len(width_demand.get(mm, []))) or 1.0
        return matched/total - 0.5*(trans/total)
    score = {k: score_from_len(k) for k in ['0_32','0_46','0_70','0_72','0_86']}

    # Сортируем по score (desc), затем по базовому порядку
    base_index = {k: i for i, k in enumerate(base)}
    priority = sorted(base, key=lambda k: (-score.get(k, 0.0), base_index[k]))
    OPT_WIDTH_PRIORITY = priority

def make_plate_name(length_m: float, width_m: float, reinforcement: str = '8п') -> str:
    """Формирует строку наименования в стиле прайса: 'Плиты ПБ 63-12-8п'.
    Для лент 0.3/0.2 записывает ширину как '0.3'/'0.2'."""
    length_dm = int(round(length_m * 10))
    if width_m < 0.5:
        width_str = '0.3' if abs(width_m - 0.3) < 1e-6 else ('0.2' if abs(width_m - 0.2) < 1e-6 else f'{width_m:.1f}'.replace('.', ','))
    else:
        width_dm = int(round(width_m * 10))
        width_str = str(width_dm)
    return f'Плиты ПБ {length_dm}-{width_str}-{reinforcement}'


def parse_name_to_sizes(name: str) -> tuple:
    """Достаёт (length_m, width_m) из строки прайса."""
    m = re.search(r'(\d+)-(\d+)', name.replace(',', '.'))
    if not m:
        return None, None
    return float(m.group(1)) / 10.0, float(m.group(2)) / 10.0


def load_price_table_from_xlsx(path: str):
    """Загружает таблицу цен вида: ключ length_dm -> {6:price,8:price,10:price,12:price}.
    Сравниваем только по длине. Ищем столбец 'Наименование' и ценовые колонки.
    Расширенные правила распознавания ценовых колонок:
      - '<x> нагрузка' (как раньше)
      - заголовки, где одновременно встречаются ('цен'|'руб'|'стоим') и цифра 6/8/10/12, например 'Цена 8п', 'Цена, руб (8)'
      - если явных колонок по нагрузкам нет, используем любой общий столбец цены ('цен'|'руб'|'стоим').
    """
    table = {}
    if pd is None:
        return table
    # Нормализуем путь: если файл не найден (возможны различия в диакритике "й"), пробуем подобрать кандидата в папке
    candidate_paths = []
    if os.path.exists(path):
        candidate_paths = [path]
    else:
        # пробуем несколько директорий: рядом с файлом, BASE_DIR и подкаталог 'банк знаний'
        search_dirs = []
        folder = os.path.dirname(path) if os.path.dirname(path) else BASE_DIR
        search_dirs.append(folder)
        search_dirs.append(BASE_DIR)
        search_dirs.append(os.path.join(BASE_DIR, 'банк знаний'))
        for d in search_dirs:
            if not os.path.isdir(d):
                continue
            for name in os.listdir(d):
                low = name.lower()
                if low.endswith('.xlsx') and ('нов' in low and 'цен' in low):
                    candidate_paths.append(os.path.join(d, name))
    if not candidate_paths:
        print('[ПРАЙС] Файл не найден. Искал около:', path)
        return table
    try:
        # Берём первый успешно распознанный файл
        chosen = None
        for p in candidate_paths:
            try:
                all_sheets = pd.read_excel(p, sheet_name=None)
                chosen = p
                break
            except Exception:
                continue
        if chosen is None:
            print('[ПРАЙС] Не удалось открыть ни один XLSX из кандидатов:', candidate_paths)
            return {}
        else:
            print('[ПРАЙС] Использую прайс-файл:', chosen)
        # Предпочтительно используем лист с датой, как на скрине: "24.06.2024"
        preferred_sheet = None
        for s in list(all_sheets.keys()):
            if '24.06.2024' in str(s):
                preferred_sheet = s
                break
        sheets_iter = [preferred_sheet] if preferred_sheet in all_sheets else list(all_sheets.keys())
        for sheet_name in sheets_iter:
            df = all_sheets[sheet_name]
            try:
                print(f"[ПРАЙС] Лист: {sheet_name} | колонки: {[str(c) for c in df.columns]}")
            except Exception:
                pass
            # Жёсткие имена колонок из вашего файла
            name_col = next((c for c in df.columns if str(c).strip().lower() == 'наименование'), None) or \
                       next((c for c in df.columns if 'наимен' in str(c).lower()), None)
            if name_col is None:
                try:
                    print('[ПРАЙС] Не найден столбец наименования на листе, пропускаю')
                except Exception:
                    pass
                continue
            # поищем столбцы цен по нагрузкам и/или общий столбец цены
            load_cols = {}
            # Жёсткая привязка к заголовкам: "6 нагрузка", "8 нагрузка", ...
            header_map = {6: None, 8: None, 10: None, 12: None}
            for c in df.columns:
                cl = str(c).strip().lower()
                if cl == '6 нагрузка':
                    header_map[6] = c
                elif cl == '8 нагрузка':
                    header_map[8] = c
                elif cl == '10 нагрузка':
                    header_map[10] = c
                elif cl == '12 нагрузка':
                    header_map[12] = c
            for k,v in header_map.items():
                if v is not None:
                    load_cols[k] = v
            # общий ценовой столбец (если он один для всех нагрузок)
            simple_price_col = next((c for c in df.columns if any(k in str(c).lower() for k in ['цен', 'руб', 'стоим'])), None)
            for c in df.columns:
                cl = str(c).lower()
                # вариант 1: "<число> нагрузка"
                m = re.search(r'(\d+)\s*нагруз', cl)
                if m:
                    load_cols[int(m.group(1))] = c
                    continue
                # вариант 2: заголовок с упоминанием цены и кода нагрузки (6/8/10/12)
                m2 = re.search(r'(?:цен|руб|стоим)[^\d]{0,10}(6|8|10|12)\b', cl)
                if not m2:
                    m2 = re.search(r'\b(6|8|10|12)[^\d]{0,10}(?:цен|руб|стоим)', cl)
                if m2:
                    try:
                        load_cols[int(m2.group(1))] = c
                    except Exception:
                        pass
            found_rows = 0
            for _, row in df.iterrows():
                name = str(row.get(name_col, '')).strip()
                if not name:
                    continue
                # В прайсе имена вида "ПБ 38-12" — достаём длину из этих двух чисел
                L, _ = parse_name_to_sizes(name)
                if L is None:
                    continue
                key = int(round(L*10))
                price_by_load = {}
                if load_cols:
                    for load_code, col in load_cols.items():
                        try:
                            val = row[col]
                            if pd.notna(val):
                                price_by_load[load_code] = float(str(val).replace(' ', '').replace(',', '.'))
                        except Exception:
                            pass
                elif simple_price_col is not None:
                    try:
                        val = row[simple_price_col]
                        if pd.notna(val):
                            price_val = float(str(val).replace(' ', '').replace(',', '.'))
                            # одинаковая цена для всех нагрузок, если нет отдельных столбцов
                            for load_code in [6, 8, 10, 12]:
                                price_by_load[load_code] = price_val
                    except Exception:
                        pass
                if price_by_load:
                    table[key] = price_by_load
                    found_rows += 1
            try:
                print(f"[ПРАЙС] Считано позиций на листе: {found_rows}")
            except Exception:
                pass
    except Exception:
        return {}
    return table


def sync_price_xlsx_to_db(xlsx_path: str = PRICE_XLSX_PATH, db_path: str = PRICE_DB_PATH,
                          sheet_hint: str = '24.06.2024') -> int:
    """Заливает прайс из XLSX в SQLite (`prices`):
    columns: length_dm INTEGER, load_code INTEGER, price REAL.
    Возвращает число записанных строк."""
    if pd is None:
        return 0
    # Загружаем словарь через уже отлаженную функцию
    price_table = load_price_table_from_xlsx(xlsx_path)
    if not price_table:
        return 0
    # Плоский список (length_dm, load, price)
    rows = []
    for length_dm, loads in price_table.items():
        for load_code, price in loads.items():
            rows.append((int(length_dm), int(load_code), float(price)))

    # Пишем в SQLite
    conn = sqlite3.connect(db_path)
    try:
        cur = conn.cursor()
        cur.execute('CREATE TABLE IF NOT EXISTS prices (length_dm INTEGER, load_code INTEGER, price REAL, PRIMARY KEY(length_dm, load_code))')
        cur.executemany('INSERT OR REPLACE INTO prices (length_dm, load_code, price) VALUES (?,?,?)', rows)
        conn.commit()
        return len(rows)
    finally:
        conn.close()


def find_price_from_db(length_m: float, load_code: int = 8, db_path: str = PRICE_DB_PATH) -> float:
    """Ищет цену в БД с допуском ±1 дм, если нет точной длины."""
    length_dm = int(round(length_m * 10))
    conn = sqlite3.connect(db_path)
    try:
        cur = conn.cursor()
        # гарантируем схему
        cur.execute('CREATE TABLE IF NOT EXISTS prices (length_dm INTEGER, load_code INTEGER, price REAL, PRIMARY KEY(length_dm, load_code))')
        cur.execute('SELECT price FROM prices WHERE length_dm=? AND load_code=?', (length_dm, load_code))
        row = cur.fetchone()
        if row:
            return float(row[0])
        # допуск ±1 дм
        cur.execute('SELECT price FROM prices WHERE ABS(length_dm-?)<=1 AND load_code=? ORDER BY ABS(length_dm-?) LIMIT 1', (length_dm, load_code, length_dm))
        row = cur.fetchone()
        return float(row[0]) if row else None
    finally:
        conn.close()


def find_price_for_plate(price_table: dict, length_m: float, load_code: int = 8) -> float:
    """Возвращает цену по длине и нагрузке (ширину игнорируем)."""
    key = int(round(length_m*10))
    if key in price_table and load_code in price_table[key]:
        return price_table[key][load_code]
    # ближайшая длина ±1 дм
    for Ldm, loads in price_table.items():
        if abs(Ldm - key) <= 1 and load_code in loads:
            return loads[load_code]
    return None


def approximate_weight_kg(length_m: float, width_m: float, thickness_m: float = 0.22) -> float:
    volume = length_m * width_m * thickness_m
    return round(volume * 2400, 1)


def optimize_cuts_pulp(orders: dict) -> list[dict]:
    """Оптимизирует раскрой стандартных плит 1200мм по ширине с помощью PuLP.

    Вход:
      orders: словарь {номинальная_ширина_мм: требуемое_количество}, например {300:4,500:3,700:2,900:2}

    Модель:
      - Есть стандартная плита шириной 1200мм. Один рез по выбранному типу даёт 2 куска: main и rest (остаток).
      - Для каждого типа реза i создаём целую переменную x_i — сколько плит резать этим способом (сколько исходных плат 1200мм).
      - Для каждого заказа w создаём переменную y_w — сколько плит нужной ширины будет выполнено.
      - Для распределения кусков вводим целочисленные переменные a_iw_main и a_iw_rest —
        сколько кусков типа main/rest от реза i пойдут на заказ ширины w.

    Ограничения:
      - y_w = sum_i (a_iw_main + a_iw_rest) и y_w >= orders[w]
      - Для каждого i: sum_w a_iw_main <= x_i  (main-кусков от i не больше, чем x_i)
      - Для каждого i: sum_w a_iw_rest <= x_i  (rest-кусков от i не больше, чем x_i)
      - a_iw_main создаются только если номинал w попадает в допустимый диапазон main у реза i (иначе переменная не создаётся)
      - a_iw_rest аналогично для диапазона rest

    Целевая функция:
      - минимизировать суммарное число резов sum_i x_i
      - плюс штраф 1000 за каждый неиспользованный остаток: для i это (x_i - sum_w a_iw_rest)
        Итого:  sum_i x_i + 1000 * sum_i (x_i - sum_w a_iw_rest)

    Возвращает список словарей с оптимальным количеством резов по каждому типу:
      [{"cut_id": "cut300", "qty": 2, "main_range": (260,320), "rest_range": (880,940)}, ...]

    Примечания:
      - Функция самодостаточна: при отсутствии pulp печатает предупреждение и возвращает пустой результат.
      - Ширины и диапазоны — в миллиметрах.
    """
    try:
        from pulp import LpProblem, LpMinimize, LpVariable, LpInteger, lpSum, LpStatusOptimal, value
    except Exception:
        print('[OPT] Модуль pulp не установлен. Пропускаю оптимизацию раскроя.')
        return []

    # Описание доступных типов резов (в мм)
    CUT_OPTIONS = [
        {"id": "cut300", "main": (260, 320), "rest": (880, 940)},
        {"id": "cut500", "main": (460, 530), "rest": (670, 740)},
        {"id": "cut700", "main": (660, 720), "rest": (480, 540)},
        {"id": "cut900", "main": (860, 920), "rest": (280, 340)},
    ]

    # Нормализация входа: ключи — int, значения — int >= 0
    orders_mm = {}
    for k, v in (orders or {}).items():
        try:
            w = int(k)
            q = int(v)
            if q > 0:
                orders_mm[w] = q
        except Exception:
            continue
    if not orders_mm:
        return []

    widths = sorted(orders_mm.keys())

    # Подготовим списки источников для каждого заказа w
    # Какие резы могут дать main-кусок для w, и какие — rest-кусок для w
    main_sources = {w: [] for w in widths}
    rest_sources = {w: [] for w in widths}
    for i, opt in enumerate(CUT_OPTIONS):
        m_lo, m_hi = opt["main"]
        r_lo, r_hi = opt["rest"]
        for w in widths:
            if m_lo <= w <= m_hi:
                main_sources[w].append(i)
            if r_lo <= w <= r_hi:
                rest_sources[w].append(i)

    # Если для какого-то w вообще нет источников — задача невыполнима, вернём пусто
    for w in widths:
        if not main_sources[w] and not rest_sources[w]:
            print(f"[OPT] Для ширины {w} мм нет допустимых резов. Задача может быть невыполнима.")

    # Модель
    prob = LpProblem('cut_optimization', LpMinimize)

    # Переменные x_i — количество исходных плит 1200мм, распиленных типом i
    x = {
        i: LpVariable(f"x_{CUT_OPTIONS[i]['id']}", lowBound=0, cat=LpInteger)
        for i in range(len(CUT_OPTIONS))
    }

    # Переменные распределения кусков по заказам: a_iw_main, a_iw_rest
    a_main = {i: {} for i in range(len(CUT_OPTIONS))}
    a_rest = {i: {} for i in range(len(CUT_OPTIONS))}
    for i, opt in enumerate(CUT_OPTIONS):
        # main
        for w in widths:
            if i in main_sources[w]:
                a_main[i][w] = LpVariable(f"a_main_{opt['id']}_{w}", lowBound=0, cat=LpInteger)
        # rest
        for w in widths:
            if i in rest_sources[w]:
                a_rest[i][w] = LpVariable(f"a_rest_{opt['id']}_{w}", lowBound=0, cat=LpInteger)

    # Переменные y_w — выполненное количество для каждого заказа
    y = {w: LpVariable(f"y_{w}", lowBound=0, cat=LpInteger) for w in widths}

    # Связи y_w с распределениями
    for w in widths:
        lhs = []
        for i in main_sources[w]:
            lhs.append(a_main[i][w])
        for i in rest_sources[w]:
            lhs.append(a_rest[i][w])
        if lhs:
            prob += (y[w] == lpSum(lhs)), f"def_y_{w}"
        else:
            # Нет источников — y_w фиксируем 0 (оставим, но требование всё равно зададим ниже)
            prob += (y[w] == 0), f"def_y_{w}"

    # Требуем обеспечить объёмы заказов
    for w in widths:
        prob += (y[w] >= orders_mm[w]), f"demand_{w}"

    # Ограничение по наличию кусков каждого типа реза: на каждый x_i есть не более x_i main и x_i rest
    for i, opt in enumerate(CUT_OPTIONS):
        if a_main[i]:
            prob += (lpSum(a_main[i].values()) <= x[i]), f"main_cap_{opt['id']}"
        else:
            # если переменных нет, то суммарно 0 <= x[i] всегда верно; ограничение опускаем
            pass
        if a_rest[i]:
            prob += (lpSum(a_rest[i].values()) <= x[i]), f"rest_cap_{opt['id']}"

    # Цель: минимизировать число резов + штраф за неиспользованный остаток
    # penalty = 1000 * sum_i (x_i - sum_w a_iw_rest)
    penalty_terms = []
    for i, opt in enumerate(CUT_OPTIONS):
        used_rest = lpSum(a_rest[i].values()) if a_rest[i] else 0
        penalty_terms.append(x[i] - used_rest)

    objective = lpSum(x.values()) + 1000 * lpSum(penalty_terms)
    prob += objective

    # Решаем
    prob.solve()

    # Проверка статуса
    try:
        status = prob.status
        # 1 — LpStatusOptimal в pulp
        if status != 1:
            print(f"[OPT] Решение не оптимально. Статус: {status}")
    except Exception:
        pass

    # Собираем результат
    result = []
    for i, opt in enumerate(CUT_OPTIONS):
        try:
            qty = int(round(x[i].value()))
        except Exception:
            qty = 0
        result.append({
            "cut_id": opt["id"],
            "qty": qty if qty > 0 else 0,
            "main_range": tuple(opt["main"]),
            "rest_range": tuple(opt["rest"]),
        })

    return result


def optimize_with_lengths(width_demand: dict[int, list[float]], pair_map: dict[int, int], trans_cut_penalty: float = 0.5) -> dict:
    """ILP оптимизация по длинам: для каждой группы ширины w
    распределяем длины из основного списка по длинам из парного списка.

    - width_demand: {w_mm: [L1, L2, ...]} списки длин для основной ширины
    - pair_map: сопоставление основной ширины w_mm -> парная ширина p_mm (например 320->880)
    - trans_cut_penalty: штраф за несовпадение длины (0..1) в "штучных" единицах, чтобы предпочитать совпадения

    Возвращает словарь с оценкой:
      {w_mm: { 'matched': k, 'trans_cuts': t, 'plan': [(L_main, L_pair, match:0/1), ...] }}
    """
    try:
        from pulp import LpProblem, LpMaximize, LpVariable, LpBinary, lpSum, value
    except Exception:
        # fallback: простое жадное сопоставление
        result = {}
        for w, main_ls in width_demand.items():
            pair_ls = width_demand.get(pair_map.get(w, -1), [])
            a = sorted(round(x, 2) for x in main_ls)
            b = sorted(round(x, 2) for x in pair_ls)
            i = j = matches = 0
            plan = []
            while i < len(a) and j < len(b):
                if abs(a[i] - b[j]) <= 0.05:
                    matches += 1; plan.append((a[i], b[j], 1)); i += 1; j += 1
                elif a[i] < b[j]:
                    plan.append((a[i], None, 0)); i += 1
                else:
                    j += 1
            while i < len(a):
                plan.append((a[i], None, 0)); i += 1
            result[w] = {'matched': matches, 'trans_cuts': max(0, len(a) - matches), 'plan': plan}
        return result

    result = {}
    for w_mm, main_ls in width_demand.items():
        p_mm = pair_map.get(w_mm)
        if p_mm is None:
            continue
        pair_ls = width_demand.get(p_mm, [])
        A = [round(x, 2) for x in main_ls]
        B = [round(x, 2) for x in pair_ls]
        n = len(A); m = len(B)
        prob = LpProblem(f"len_match_{w_mm}", LpMaximize)
        x = [[LpVariable(f"x_{i}_{j}", lowBound=0, upBound=1, cat=LpBinary) for j in range(m)] for i in range(n)]
        y = [LpVariable(f"y_{i}", lowBound=0, upBound=1, cat=LpBinary) for i in range(n)]  # 1 если есть парный
        # Связи: каждому i максимум один j
        for i in range(n):
            prob += lpSum(x[i][j] for j in range(m)) == y[i]
        # Каждый j максимум к одному i
        for j in range(m):
            prob += lpSum(x[i][j] for i in range(n)) <= 1
        # Цель: максимизировать совпадения по длине, штрафуя несовпадения как trans cut
        # score = sum_i sum_j x_ij * match(i,j) - trans_cut_penalty * sum_i (y_i - best_match)
        # Реализуем: match(i,j)=1 если |A[i]-B[j]|<=0.05 else 1 - penalty
        from math import fabs
        match = [[1.0 if fabs(A[i]-B[j]) <= 0.05 else 1.0 - trans_cut_penalty for j in range(m)] for i in range(n)]
        prob += lpSum(match[i][j] * x[i][j] for i in range(n) for j in range(m))
        prob.solve()
        plan = []
        matched = 0
        used_j = set()
        for i in range(n):
            paired = False
            for j in range(m):
                try:
                    if value(x[i][j]) >= 0.5:
                        paired = True
                        used_j.add(j)
                        good = 1 if abs(A[i]-B[j]) <= 0.05 else 0
                        matched += good
                        plan.append((A[i], B[j], good))
                        break
                except Exception:
                    continue
            if not paired:
                plan.append((A[i], None, 0))
        trans = sum(1 for a, b, good in plan if b is not None and good == 0)
        result[w_mm] = {'matched': matched, 'trans_cuts': trans, 'plan': plan}
    return result

def load_cut_price_from_docx(path: str) -> float:
    """Пытается извлечь цену продольного реза из DOCX. Возвращает 0, если не удалось."""
    if Document is None or not os.path.exists(path):
        return 0.0
    try:
        doc = Document(path)
        text = '\n'.join([p.text for p in doc.paragraphs])
        # Находим фразы про рез/вдоль/продольн и число рядом
        candidates = []
        for m in re.finditer(r'(рез|вдоль|продоль)[^\d]{0,20}(\d+[\s\u202f\,\.]?\d*)', text.lower()):
            try:
                val = float(m.group(2).replace(' ', '').replace('\u202f', '').replace(',', '.'))
                candidates.append(val)
            except Exception:
                pass
        if candidates:
            return float(max(candidates))
        # также ищем в таблицах
        for table in doc.tables:
            for row in table.rows:
                row_text = ' '.join(c.text for c in row.cells).lower()
                if any(k in row_text for k in ['рез', 'вдоль', 'продоль']):
                    nums = re.findall(r'\d+[\s\u202f\,\.]?\d*', row_text)
                    for s in nums:
                        try:
                            val = float(s.replace(' ', '').replace('\u202f', '').replace(',', '.'))
                            candidates.append(val)
                        except Exception:
                            pass
        return float(max(candidates)) if candidates else 0.0
    except Exception:
        return 0.0


def build_procurement_items():
    """Формирует реальные позиции закупки с учётом назначения реза.
    Если есть OPT_PLAN — используем его, иначе fallback на PLATES_*.
    Возвращает список dict: {length, width, qty, long_cuts, trans_cuts}.
    """
    global OPT_PLAN
    items = []
    
    # Если оптимизатор дал план — используем его
    if OPT_PLAN and OPT_PLAN.get('actions'):
        for act in OPT_PLAN['actions']:
            src_type, W1, W2, L, qty, lc, tc = act
            W1_m = W1 / 1000.0; W2_m = W2 / 1000.0 if W2 else 0
            if src_type == 'split':
                # split: закупаем исходную плиту 1.2×L с продольным резом
                # Она даст две полосы: W1 и W2, но в смете — одна позиция 1.2м
                items.append({'length': round(L, 2), 'width': 1.2, 'qty': qty, 'long_cuts': lc, 'trans_cuts': tc, 'purpose': 'split_source'})
            elif src_type == 'narrow':
                # narrowing: закупаем плиту шириной W2 (исходная) и режем до W1
                # В смете показываем W2 (исходную ширину)
                items.append({'length': round(L, 2), 'width': W2_m, 'qty': qty, 'long_cuts': lc, 'trans_cuts': tc, 'purpose': 'narrow_source'})
            elif src_type == 'solid':
                items.append({'length': round(L, 2), 'width': W1_m, 'qty': qty, 'long_cuts': lc, 'trans_cuts': tc, 'purpose': 'solid'})
        # агрегируем
        agg = {}
        for it in items:
            key = (it['length'], it['width'], it['long_cuts'], it['trans_cuts'])
            agg[key] = agg.get(key, 0) + it['qty']
        result = []
        for (L, W, long_cuts, trans_cuts), qty in sorted(agg.items(), key=lambda x: (x[0][1], x[0][0])):
            result.append({'length': L, 'width': W, 'qty': qty, 'long_cuts': long_cuts, 'trans_cuts': trans_cuts})
        return result
    
    # Fallback: старая логика с PLATES_*
    def mismatch_count(main_list: list[float], pair_demand: list[float]) -> int:
        if not main_list or not pair_demand:
            return 0
        a = sorted(round(x, 2) for x in main_list)
        b = sorted(round(x, 2) for x in pair_demand)
        i = j = matches = 0
        while i < len(a) and j < len(b):
            if abs(a[i] - b[j]) <= 0.05:
                matches += 1; i += 1; j += 1
            elif a[i] < b[j]:
                i += 1
            else:
                j += 1
        return max(0, min(len(main_list), len(pair_demand)) - matches)

    pair_plan = {
        '0.32': mismatch_count(PLATES_0_32, PLATES_0_88),
        '0.46': mismatch_count(PLATES_0_46, PLATES_0_74),
        '0.72': mismatch_count(PLATES_0_72, PLATES_0_48),
        '0.70': mismatch_count(PLATES_0_70, PLATES_0_50),
        '0.86': mismatch_count(PLATES_0_86, PLATES_0_34),
    }
    # 1.2 без реза
    for L in PLATES_1_2:
        items.append({'length': round(L, 1), 'width': 1.2, 'qty': 1, 'long_cuts': 0, 'trans_cuts': 0, 'purpose': 'as_is'})
    # 1.5 -> 1.2 + 0.3: две позиции
    #   - ПБ L-12 (без реза)
    #   - Лента L-0.3 (как отдельная позиция, с продольным резом 1 шт)
    for L in PLATES_1_5_TO_1_2:
        items.append({'length': round(L, 1), 'width': 1.2, 'qty': 1, 'long_cuts': 0, 'trans_cuts': 0, 'purpose': 'to_1_2_main'})
        items.append({'length': round(L, 1), 'width': 0.3, 'qty': 1, 'long_cuts': 1, 'trans_cuts': 0, 'purpose': 'to_1_2_strip'})
    # 1.2 -> 1.0 + 0.2: две позиции (вернули прежнее поведение)
    #   - ПБ L-10-8п — с продольным резом
    #   - Лента L-0.2-8п — как отдельная позиция, тарифицируется и учитывает 1 продольный рез
    for L in PLATES_1_0:
        items.append({'length': round(L, 1), 'width': 1.0, 'qty': 1, 'long_cuts': 1, 'trans_cuts': 0, 'purpose': 'to_1_0_main'})
        items.append({'length': round(L, 1), 'width': 0.2, 'qty': 1, 'long_cuts': 1, 'trans_cuts': 0, 'purpose': 'to_1_0_strip'})
    
    # Плиты с меньшей шириной (получаются резом из 1.2м)
    # 1.2 -> 1.08 + 0.12
    for L in PLATES_1_08:
        items.append({'length': round(L, 1), 'width': 1.08, 'qty': 1, 'long_cuts': 1, 'trans_cuts': 0, 'purpose': 'to_1_08_main'})
        items.append({'length': round(L, 1), 'width': 0.12, 'qty': 1, 'long_cuts': 1, 'trans_cuts': 0, 'purpose': 'to_1_08_strip'})
    
    # 1.2 -> 0.46 + 0.74
    for L in PLATES_0_46:
        items.append({'length': round(L, 1), 'width': 0.46, 'qty': 1, 'long_cuts': 1, 'trans_cuts': 0, 'purpose': 'to_0_46_main'})
        items.append({'length': round(L, 1), 'width': 0.74, 'qty': 1, 'long_cuts': 1, 'trans_cuts': 0, 'purpose': 'to_0_46_strip'})
    
    # 1.2 -> 0.32 + 0.88
    # 1.2 -> 0.32 + 0.88 (часть 0.88 может требовать поперечного реза если длина отличается)
    mismatch = pair_plan['0.32']
    for idx, L in enumerate(PLATES_0_32):
        items.append({'length': round(L, 1), 'width': 0.32, 'qty': 1, 'long_cuts': 1, 'trans_cuts': 0, 'purpose': 'to_0_32_main'})
        trans = 1 if idx < mismatch else 0
        items.append({'length': round(L, 1), 'width': 0.88, 'qty': 1, 'long_cuts': 1, 'trans_cuts': trans, 'purpose': 'to_0_32_strip'})
    
    # 1.2 -> 0.72 + 0.48
    mismatch = pair_plan['0.72']
    for idx, L in enumerate(PLATES_0_72):
        items.append({'length': round(L, 1), 'width': 0.72, 'qty': 1, 'long_cuts': 1, 'trans_cuts': 0, 'purpose': 'to_0_72_main'})
        trans = 1 if idx < mismatch else 0
        items.append({'length': round(L, 1), 'width': 0.48, 'qty': 1, 'long_cuts': 1, 'trans_cuts': trans, 'purpose': 'to_0_72_strip'})
    
    # 1.2 -> 0.70 + 0.50
    mismatch = pair_plan['0.70']
    for idx, L in enumerate(PLATES_0_70):
        items.append({'length': round(L, 1), 'width': 0.70, 'qty': 1, 'long_cuts': 1, 'trans_cuts': 0, 'purpose': 'to_0_70_main'})
        trans = 1 if idx < mismatch else 0
        items.append({'length': round(L, 1), 'width': 0.50, 'qty': 1, 'long_cuts': 1, 'trans_cuts': trans, 'purpose': 'to_0_70_strip'})
    
    # 1.2 -> 0.86 + 0.34
    mismatch = pair_plan['0.86']
    for idx, L in enumerate(PLATES_0_86):
        items.append({'length': round(L, 1), 'width': 0.86, 'qty': 1, 'long_cuts': 1, 'trans_cuts': 0, 'purpose': 'to_0_86_main'})
        trans = 1 if idx < mismatch else 0
        items.append({'length': round(L, 1), 'width': 0.34, 'qty': 1, 'long_cuts': 1, 'trans_cuts': trans, 'purpose': 'to_0_86_strip'})
    # агрегируем по (L,W,cuts)
    agg = {}
    for it in items:
        key = (it['length'], it['width'], it['long_cuts'], it['trans_cuts'])
        agg[key] = agg.get(key, 0) + it['qty']
    result = []
    for (L, W, long_cuts, trans_cuts), qty in sorted(agg.items(), key=lambda x: (x[0][1], x[0][0])):
        result.append({'length': L, 'width': W, 'qty': qty, 'long_cuts': long_cuts, 'trans_cuts': trans_cuts})
    return result


def build_price_rows(price_table: dict, reinforcement_code: int = 8):
    """Формирует строки сметы на базе реальных позиций закупки.
    Цена позиции = цена плиты (по длине/ширине и нагрузке) + цена резов на плиту.
    Возвращает (rows, total_sum)."""
    items = build_procurement_items()
    rows = []
    total = 0.0
    idx = 1
    for it in items:
        L, W, qty = it['length'], it['width'], it['qty']
        long_cuts, trans_cuts = it['long_cuts'], it['trans_cuts']
        name = make_plate_name(L, W)
        # 1) пытаемся взять из БД, если есть прайс-таблица
        # Используем БД (цены полностью перенесены)
        # Для плит с меньшей шириной используем нагрузку 6, для стандартных - 8
        load_code = 6 if W < 1.0 else reinforcement_code
        db_price = get_price(L, load_code, PRICE_DB_PATH)
        # 2) fallback — из XLSX-таблицы, если БД пустая
        base_price_1_2m = db_price if db_price is not None else (find_price_for_plate(price_table, L, load_code) or 0.0)
        
        # 3) Корректируем цену пропорционально ширине плиты
        # Цены в БД даны для плит шириной 1.2м, корректируем для других ширин
        if base_price_1_2m > 0:
            # Рассчитываем цену пропорционально ширине
            width_factor = W / 1.2  # коэффициент пропорциональности
            base_price = base_price_1_2m * width_factor
        else:
            base_price = 0.0
        cuts_cost = long_cuts * (LONG_CUT_PRICE_PER_M * L) + trans_cuts * TRANSVERSE_CUT_PRICE
        unit_price = base_price + cuts_cost
        weight = approximate_weight_kg(L, W)
        row_sum = unit_price * qty
        total += row_sum
        rows.append([
            idx,
            name,
            qty,
            'шт',
            f'{weight:.0f}',
            f'{unit_price:,.2f}'.replace(',', ' ').replace('.', ','),
            f'{row_sum:,.2f}'.replace(',', ' ').replace('.', ',')
        ])
        idx += 1
    return rows, total


# ---------- Рендер ----------

def _draw_segment(ax, x0: float, length: float, color: str, label: str, y: float = 0.0, height: float = TRACK_WIDTH_M):
    rect = patches.Rectangle((x0, y), length, height, linewidth=1, edgecolor='black', facecolor=color, alpha=0.85)
    ax.add_patch(rect)
    ax.text(x0 + length/2, y + height/2, label, ha='center', va='center', fontsize=8, color='white', weight='bold')


def _draw_split_plate(ax, x0: float, length: float, main_w: float, rest_w: float, label_main: str, label_rest: str | None = None):
    # Основа плиты 1.2
    rect = patches.Rectangle((x0, 0.0), length, TRACK_WIDTH_M, linewidth=1.2, edgecolor='black', facecolor='#ecf0f1', alpha=1.0)
    ax.add_patch(rect)
    # Полоса основной ширины (снизу вверх)
    main_rect = patches.Rectangle((x0, 0.0), length, main_w, linewidth=0.8, edgecolor='black', facecolor='#2ecc71', alpha=0.9)
    ax.add_patch(main_rect)
    # Разделительная линия реза
    ax.plot([x0, x0 + length], [main_w, main_w], color='black', linestyle='--', linewidth=1)
    # Метки
    ax.text(x0 + length/2, main_w/2, label_main, ha='center', va='center', fontsize=8, color='white', weight='bold')
    if label_rest and rest_w > 0.02:
        ax.text(x0 + length/2, main_w + rest_w/2, label_rest, ha='center', va='center', fontsize=7, color='#2c3e50')


def build_layout_sequence():
    """Формирует последовательность сегментов вдоль дорожки.
    Если есть OPT_PLAN — используем его, иначе fallback на PLATES_*.
    """
    global OPT_PLAN
    sequence = []

    def plate_label(L: float, W: float) -> str:
        Ldm = int(round(L * 10))
        Wdm_val = round(W * 10, 1)
        if abs(Wdm_val - int(Wdm_val)) < 1e-6:
            Wdm = str(int(Wdm_val))
        else:
            Wdm = str(Wdm_val).replace('.', ',')
        return f'ПБ {Ldm}-{Wdm}-8п'
    
    # Если есть OPT_PLAN — визуализируем по нему
    if OPT_PLAN and OPT_PLAN.get('actions'):
        for act in OPT_PLAN['actions']:
            src_type, W1, W2, L, qty, lc, tc = act
            W1_m = W1 / 1000.0; W2_m = W2 / 1000.0 if W2 else 0
            for _ in range(qty):
                if src_type == 'solid':
                    sequence.append({'length': L, 'mode': 'solid', 'label': plate_label(L, W1_m)})
                elif src_type == 'split':
                    rest_w = W2_m if W2_m < W1_m else (1.2 - W1_m)
                    rest_label = f'+{rest_w:.2f}'.replace('.', ',')
                    sequence.append({'length': L, 'mode': 'split', 'main_w': W1_m, 'rest_w': rest_w,
                                     'label_main': plate_label(L, W1_m), 'label_rest': rest_label})
                elif src_type == 'narrow':
                    # narrowing: показываем целевую W1 и отход (W2-W1)
                    delta = abs(W2_m - W1_m) if W2_m else 0
                    rest_label = f'-{delta:.2f}'.replace('.', ',') if delta > 0.001 else ''
                    sequence.append({'length': L, 'mode': 'split', 'main_w': W1_m, 'rest_w': delta,
                                     'label_main': plate_label(L, W1_m), 'label_rest': rest_label})
        return sequence
    
    # Fallback: старая логика с PLATES_*
    for L in PLATES_1_2:
        sequence.append({'length': L, 'mode': 'solid', 'label': plate_label(L, 1.2)})

    for L in PLATES_1_5_TO_1_2:
        sequence.append({'length': L, 'mode': 'solid', 'label': plate_label(L, 1.2)})

    for L in PLATES_1_0:
        sequence.append({'length': L, 'mode': 'split', 'main_w': 1.0, 'rest_w': 0.2,
                         'label_main': plate_label(L, 1.0), 'label_rest': '+0,2'})

    for L in globals().get('PLATES_1_08', []):
        sequence.append({'length': L, 'mode': 'split', 'main_w': 1.08, 'rest_w': 0.12,
                         'label_main': plate_label(L, 1.08), 'label_rest': '+0,12'})

    # Группы < 1.2 м будем добавлять в порядке приоритета OPT_WIDTH_PRIORITY
    groups_map = {
        '0_32': (globals().get('PLATES_0_32', []), 0.32, 0.88, '+0,88'),
        '0_46': (globals().get('PLATES_0_46', []), 0.46, 0.74, '+0,74'),
        '0_70': (globals().get('PLATES_0_70', []), 0.70, 0.50, '+0,50'),
        '0_72': (globals().get('PLATES_0_72', []), 0.72, 0.48, '+0,48'),
        '0_86': (globals().get('PLATES_0_86', []), 0.86, 0.34, '+0,34'),
    }
    # Если пользователь заказал пары (0.74/0.88/0.48/0.50/0.34), добавим их как отдельные “main”
    if len(globals().get('PLATES_0_74', [])):
        groups_map['0_74'] = (globals().get('PLATES_0_74', []), 0.74, 0.46, '+0,46')
    if len(globals().get('PLATES_0_88', [])):
        groups_map['0_88'] = (globals().get('PLATES_0_88', []), 0.88, 0.32, '+0,32')
    if len(globals().get('PLATES_0_48', [])):
        groups_map['0_48'] = (globals().get('PLATES_0_48', []), 0.48, 0.72, '+0,72')
    if len(globals().get('PLATES_0_50', [])):
        groups_map['0_50'] = (globals().get('PLATES_0_50', []), 0.50, 0.70, '+0,70')
    if len(globals().get('PLATES_0_34', [])):
        groups_map['0_34'] = (globals().get('PLATES_0_34', []), 0.34, 0.86, '+0,86')
    order = OPT_WIDTH_PRIORITY or list(groups_map.keys())
    for key in order:
        items, main_w, rest_w, rest_label = groups_map[key]
        for L in items:
            sequence.append({'length': L, 'mode': 'split', 'main_w': main_w, 'rest_w': rest_w,
                             'label_main': plate_label(L, main_w), 'label_rest': rest_label})

    return sequence


def visualize_plan(output_dir: str = 'Визуализация_Раскладки'):
    # Демонстрационный вызов оптимизации раскроя (не влияет на визуализацию)
    try:
        optimized = optimize_cuts_pulp({300: 4, 500: 3, 700: 2, 900: 2})
        print("Оптимальные резы:", optimized)
    except Exception as e:
        print("[OPT] Ошибка при оптимизации:", e)

    os.makedirs(output_dir, exist_ok=True)

    # Загружаем прайс
    price_table = load_price_table_from_xlsx(PRICE_XLSX_PATH)
    # Полная миграция цен в БД (однократно при запуске)
    try:
        init_schema(PRICE_DB_PATH)
        written = import_from_xlsx(PRICE_XLSX_PATH, PRICE_DB_PATH)
        if written:
            print(f'[ПРАЙС->БД] записано строк: {written}')
    except Exception:
        pass
    # Не читаем DOCX: используем константы LONG_CUT_PRICE_PER_M/TRANSVERSE_CUT_PRICE

    # Строим смету
    price_rows, total_sum = build_price_rows(price_table)

    seq = build_layout_sequence()
    total_length = sum(s['length'] for s in seq)

    # Настройка фигуры: 4 строки — дорожка, сводка, таблица ведомости, таблица сметы
    fig = plt.figure(figsize=(22, 14))
    gs = fig.add_gridspec(4, 1, height_ratios=[3.0, 1.0, 1.4, 1.8])
    ax_track = fig.add_subplot(gs[0, 0])
    ax_strips = fig.add_subplot(gs[1, 0])
    ax_table = fig.add_subplot(gs[2, 0])
    ax_price = fig.add_subplot(gs[3, 0])
    fig.suptitle('КЗ: Дорожка 1 (ширина 1.2 м) — раскладка, резы, ведомости и смета', fontsize=16, fontweight='bold')

    # Дорожка
    ax_track.set_xlim(0, max(total_length + 2, TRACK_LENGTH_M))
    ax_track.set_ylim(0, TRACK_WIDTH_M + 0.8)
    ax_track.set_aspect('auto')
    ax_track.spines['top'].set_visible(False)
    ax_track.spines['right'].set_visible(False)
    ax_track.set_yticks([0, 0.6, 1.2])
    ax_track.set_yticklabels(['0', '0.6', '1.2 м'])

    # Горизонтальная линейка и сетка по метрам
    ax_track.set_xlabel('Длина (м)')
    ax_track.set_xticks(range(0, int(max(total_length, TRACK_LENGTH_M)) + 1, 5))
    ax_track.grid(axis='x', linestyle=':', linewidth=0.5, alpha=0.5)

    # Рисуем границу дорожки
    track_rect = patches.Rectangle((0, 0), TRACK_LENGTH_M, TRACK_WIDTH_M, linewidth=2, edgecolor='black', facecolor='none', linestyle='--')
    ax_track.add_patch(track_rect)

    # Цвета типов
    # Цвета больше не нужны для разных типов, так как мы рисуем рез внутри 1.2

    # Рисуем последовательность
    x = 0.0
    for item in seq:
        if item.get('mode') == 'solid':
            _draw_segment(ax_track, x, item['length'], '#2ecc71', item['label'])
        else:
            _draw_split_plate(
                ax_track, x, item['length'],
                main_w=item['main_w'], rest_w=item['rest_w'],
                label_main=item['label_main'], label_rest=item.get('label_rest')
            )
        x += item['length']

    # Легенда
    legend_patches = [
        patches.Patch(facecolor='#2ecc71', edgecolor='black', label='Плита 1.2 м'),
        patches.Patch(facecolor='#ecf0f1', edgecolor='black', label='Зона реза (контур)'),
    ]
    ax_track.legend(handles=legend_patches, loc='upper right')

    # Панель сводки
    ax_strips.set_xlim(0, 100)
    ax_strips.set_ylim(0, 1)
    ax_strips.axis('off')

    txt = (
        f"Длина по плану: {total_length:.1f} м  |  Продольных резов: {LONGITUDINAL_CUTS}  |  Подрезов по длине: {LENGTH_TRIMS}\n"
        f"Остатки лент 0.3: {UNUSED_STRIPS_0_3_M_TOTAL:.1f} пог.м  |  Обрезки 0.2: {SCRAP_STRIPS_0_2_M_TOTAL:.1f} пог.м (≈ {WASTE_AREA_M2:.2f} м²)"
    )
    ax_strips.text(0.02, 0.6, txt, ha='left', va='center', fontsize=12,
                   bbox=dict(boxstyle='round,pad=0.6', facecolor='#f8f9fa', edgecolor='#bdc3c7'))

    leftovers = (
        "Остатки/обрезки:\n"
        "Ленты 0.3: 3×3.8 м; 1×2.9 м\n"
        f"Ленты 0.2 (обрезки): {', '.join(f'{L:.1f} м' for L in PLATES_1_0)}"
    )
    ax_strips.text(0.02, 0.15, leftovers, ha='left', va='center', fontsize=11,
                   bbox=dict(boxstyle='round,pad=0.5', facecolor='#eef7ff', edgecolor='#a3c9ff'))

    # Таблица: ведомость (заказ vs использовано/обрезки)
    ax_table.axis('off')

    order_list = [
        'Заказ 1.5: 3.8×3; 2.9×1',
        'Заказ 1.2: 6.3×2; 3.8×2',
        'Заказ 1.0: 6.3×1; 5.3×2; 3.8×1; 2.8×1',
    ]

    used_list = [
        '1.2 без реза: 6.3×2; 3.8×2',
        '1.5→1.2: 3.8×3; 2.9×1 (остаток 0.3)',
        '1.2→1.0: 6.3×1; 5.3×2; 3.8×1; 2.8×1 (остаток 0.2 → обрезки)',
        f'Резы: продольных {LONGITUDINAL_CUTS}; подрезов {LENGTH_TRIMS}; обрезки 0.2 ≈ {WASTE_AREA_M2:.2f} м²',
    ]

    rows = max(len(order_list), len(used_list))
    table_rows = []
    for i in range(rows):
        left = order_list[i] if i < len(order_list) else ''
        right = used_list[i] if i < len(used_list) else ''
        table_rows.append([left, right])

    col_labels = ['Список плит по заказу', 'Использовано (с учётом резов) / остатки / обрезки']

    table = ax_table.table(cellText=table_rows, colLabels=col_labels, loc='center', cellLoc='left', colLoc='center')
    table.auto_set_font_size(False)
    table.set_fontsize(11)
    table.scale(1, 1.5)

    # Таблица: смета (как на фото)
    ax_price.axis('off')
    price_headers = ['№', 'Наименование', 'Кол-во', 'Ед.', 'Вес(кг)', 'Цена', 'Сумма']
    price_table = ax_price.table(cellText=price_rows, colLabels=price_headers, loc='center', cellLoc='center', colLoc='center')
    price_table.auto_set_font_size(False)
    price_table.set_fontsize(10)
    price_table.scale(1, 1.4)
    # Диагностика: если какие-то цены не найдены (0), подсказка в заголовке
    not_priced = any(row[5].strip().startswith('0') for row in price_rows)
    title = f'Итоговая стоимость: {total_sum:,.2f} ₽'.replace(',', ' ').replace('.', ',')
    if not_priced:
        title += ' (внимание: не найдены цены для некоторых позиций — проверьте прайс)'
    ax_price.set_title(title, fontsize=12, pad=10)

    # Экспорт таблиц
    timestamp = datetime.now().strftime('%Y%m%d_%H%M')
    csv_path = os.path.join(output_dir, f'Ведомость_Дорожка_1_{timestamp}.csv')
    with open(csv_path, 'w', encoding='utf-8') as f:
        f.write('Список плит по заказу;Использовано (с учётом резов) / остатки / обрезки\n')
        for left, right in table_rows:
            f.write(f'{left};{right}\n')

    if pd is not None:
        try:
            df_v = pd.DataFrame(table_rows, columns=col_labels)
            xlsx_path_v = os.path.join(output_dir, f'Ведомость_Дорожка_1_{timestamp}.xlsx')
            df_v.to_excel(xlsx_path_v, index=False)

            df_p = pd.DataFrame(price_rows, columns=price_headers)
            xlsx_path_p = os.path.join(output_dir, f'Смета_Дорожка_1_{timestamp}.xlsx')
            with pd.ExcelWriter(xlsx_path_p, engine='openpyxl') as writer:
                df_p.to_excel(writer, index=False, sheet_name='Смета')
                df_v.to_excel(writer, index=False, sheet_name='Ведомость')
        except Exception:
            pass

    # Сохранение изображений
    png_path = os.path.join(output_dir, f'Схема_Дорожка_1_КЗ_{timestamp}.png')
    pdf_path = os.path.join(output_dir, f'Схема_Дорожка_1_КЗ_{timestamp}.pdf')
    plt.tight_layout(rect=[0, 0, 1, 0.95])
    fig.savefig(png_path, dpi=300)
    fig.savefig(pdf_path)
    plt.close(fig)

    print('[ГОТОВО] Визуализация и файлы сохранены:')
    print('  PNG:', png_path)
    print('  PDF:', pdf_path)
    print('  CSV:', csv_path)
    if pd is not None:
        print('  XLSX (ведомость):', os.path.join(output_dir, f'Ведомость_Дорожка_1_{timestamp}.xlsx'))
        print('  XLSX (смета):', os.path.join(output_dir, f'Смета_Дорожка_1_{timestamp}.xlsx'))
    return png_path, pdf_path


if __name__ == '__main__':
    visualize_plan()
